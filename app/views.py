from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth import authenticate, login, logout
from django.contrib.auth.decorators import login_required
from django.db.models import Q
from django.contrib.auth.models import User
from django.http import HttpResponse, JsonResponse, HttpResponseNotFound
from django.template.loader import render_to_string
from django.core.mail import send_mail
from django.conf import settings
from app.decorators import check_page_enabled
import time
import csv
import io
import os
from .forms import (
    SearchForm, ContactForm, UserRegistrationForm, ArticleForm, ArticleAuthorForm,
    JournalForm, JournalEditorForm, ProjectForm, ProjectContributorForm,
    MembershipRequestForm, DirectoryApplicationForm, HallOfFameApplicationForm,
    PlagiarismCheckForm, PlagiarismWorkForm, ThesisToArticleForm, ThesisToBookForm,
    ThesisToBookChapterForm, PowerPointPreparationForm, NewsArticleForm
)
from .models import (
    UserProfile, Article, ArticleAuthor, Journal, JournalEditor, Project, ProjectContributor, ProjectPayment,
    MembershipRequest, DirectoryApplication, HallOfFameApplication, PlagiarismCheck,
    PlagiarismWork, ThesisToArticle, ThesisToBook, ThesisToBookChapter, PowerPointPreparation,
    NewsTag, NewsWriter, NewsArticle, NewsComment, NewsBookmark
)

def landing(request):
    """Landing page view"""
    from app.models import SiteSettings
    site_settings = SiteSettings.get_settings()
    
    # Check if landing page is enabled
    if not site_settings.enable_landing_page:
        from django.http import HttpResponseNotFound
        return HttpResponseNotFound("Page not found")
    
    search_form = SearchForm()
    # Get latest articles for the landing page
    latest_articles = Article.objects.filter(status='approved').order_by('-created_at')[:6]
    return render(request, 'app/landing.html', {
        'search_form': search_form,
        'latest_articles': latest_articles,
        'site_settings': site_settings,
    })

def browse(request):
    """Browse articles view"""
    from django.db.models import Q
    
    # Get search query from GET parameter or form
    search_query = request.GET.get('search', '').strip()
    search_form = SearchForm(request.GET or None)
    
    # If search parameter is provided, use it
    if search_query:
        search_form = SearchForm({'query': search_query})
        results = Article.objects.filter(
            Q(title__icontains=search_query) | Q(abstract__icontains=search_query)
        )
    elif search_form.is_valid():
        query = search_form.cleaned_data['query']
        # Search articles in database
        results = Article.objects.filter(
            Q(title__icontains=query) | Q(abstract__icontains=query)
        )
    else:
        results = Article.objects.none()
    
    return render(request, 'app/browse.html', {
        'search_form': search_form,
        'results': results,
    })

def submit(request):
    """Submit research view"""
    contact_form = ContactForm(request.POST or None)
    
    if request.method == 'POST':
        if contact_form.is_valid():
            # Here you could save to a Contact model if needed
            messages.success(request, 'Thank you for your submission!')
            return redirect('app:landing')
    
    return render(request, 'app/submit.html', {
        'contact_form': contact_form,
    })

def about(request):
    """About page view"""
    return render(request, 'app/about.html')

def auth(request):
    """Login/Register page view"""
    if request.method == 'POST':
        form_type = request.POST.get('form_type')
        
        if form_type == 'login':
            username_or_email = request.POST.get('username_or_email')
            password = request.POST.get('password')
            
            # Try to authenticate with username or email
            user = authenticate(request, username=username_or_email, password=password)
            if user is None:
                # Try with email
                try:
                    user_obj = User.objects.get(email=username_or_email)
                    user = authenticate(request, username=user_obj.username, password=password)
                except User.DoesNotExist:
                    pass
            
            if user is not None:
                login(request, user)
                messages.success(request, f'Welcome back, {user.username}!')
                if user.is_superuser or user.is_staff:
                    return redirect('app:dashboard')
                else:
                    return redirect('app:user_dashboard')
            else:
                messages.error(request, 'Invalid username/email or password')
            
        elif form_type == 'register':
            # Create username from email if not provided
            email = request.POST.get('email')
            name = request.POST.get('name', '')
            username = request.POST.get('username', email.split('@')[0] if email else 'user')
            
            # Check if email is already registered
            if email and User.objects.filter(email=email).exists():
                messages.error(request, 'A user with this email already exists. Please use a different email address or try logging in.')
                return render(request, 'app/auth.html')
            
            # Ensure username is unique
            base_username = username
            counter = 1
            while User.objects.filter(username=username).exists():
                username = f"{base_username}{counter}"
                counter += 1
            
            # Split name into first_name and last_name
            name_parts = name.split(' ', 1) if name else ['', '']
            first_name = name_parts[0] if len(name_parts) > 0 else ''
            last_name = name_parts[1] if len(name_parts) > 1 else ''
            
            registration_data = {
                'username': username,
                'email': email,
                'password1': request.POST.get('password'),
                'password2': request.POST.get('confirm_password'),
            }
            
            form = UserRegistrationForm(registration_data)
            
            if form.is_valid():
                user = form.save()
                # Set first_name and last_name
                user.first_name = first_name
                user.last_name = last_name
                user.save()
                
                # Create user profile with additional data
                UserProfile.objects.create(
                    user=user,
                    phone=request.POST.get('phone', ''),
                    country=request.POST.get('country', ''),
                    profile_photo=request.FILES.get('profile_photo')
                )
                # Automatically log in the user after registration
                login(request, user)
                messages.success(request, f'Registration successful! Welcome, {user.username}!')
                if user.is_superuser or user.is_staff:
                    return redirect('app:dashboard')
                else:
                    return redirect('app:user_dashboard')
            else:
                # Display form errors
                for field, errors in form.errors.items():
                    for error in errors:
                        messages.error(request, f'{field}: {error}')
                # If form is invalid, stay on auth page to show errors
                return render(request, 'app/auth.html')
    
    return render(request, 'app/auth.html')

def logout_view(request):
    """Logout view"""
    logout(request)
    messages.success(request, 'You have been logged out successfully.')
    return redirect('app:landing')

@check_page_enabled('enable_indexed_articles_page')
def indexed_articles(request):
    """Indexed Articles page view"""
    from django.utils import timezone
    from datetime import timedelta
    from django.db.models import Count
    from django.core.paginator import Paginator
    
    # Show both approved and pending articles
    articles = Article.objects.filter(status__in=['approved', 'pending']).order_by('-created_at')
    
    # Get unique disciplines (subjects) with counts
    subjects = Article.objects.filter(status__in=['approved', 'pending']).values('discipline').annotate(count=Count('id')).order_by('discipline')
    subjects_list = [{'name': s['discipline'], 'count': s['count']} for s in subjects if s['discipline']]
    
    # Get unique journals with counts
    journals = Article.objects.filter(status__in=['approved', 'pending']).exclude(journal_name='').values('journal_name').annotate(count=Count('id')).order_by('journal_name')
    journals_list = [{'name': j['journal_name'], 'count': j['count']} for j in journals if j['journal_name']]
    
    # Get unique years with counts
    years = Article.objects.filter(status__in=['approved', 'pending']).exclude(year_of_publication__isnull=True).values('year_of_publication').annotate(count=Count('id')).order_by('-year_of_publication')
    years_list = [{'name': str(y['year_of_publication']), 'count': y['count']} for y in years if y['year_of_publication']]
    
    # Pagination
    try:
        per_page = int(request.GET.get('limit', 40))  # Default 40 items per page
        if per_page not in [20, 40, 60, 100]:
            per_page = 40
    except (ValueError, TypeError):
        per_page = 40
    paginator = Paginator(articles, per_page)
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)
    
    # Calculate record range
    start_record = (page_obj.number - 1) * paginator.per_page + 1
    end_record = min(start_record + paginator.per_page - 1, paginator.count)
    
    # Add time ago information for each article
    articles_with_time = []
    for article in page_obj:
        if article.created_at:
            time_diff = timezone.now() - article.created_at
            if time_diff.days < 1:
                time_ago = "Today"
            elif time_diff.days == 1:
                time_ago = "1 day ago"
            elif time_diff.days < 30:
                time_ago = f"{time_diff.days} days ago"
            elif time_diff.days < 60:
                time_ago = "1 month ago"
            elif time_diff.days < 365:
                months = time_diff.days // 30
                time_ago = f"{months} months ago"
            else:
                years = time_diff.days // 365
                time_ago = f"{years} year{'s' if years > 1 else ''} ago"
        else:
            time_ago = "Recently"
        
        articles_with_time.append({
            'article': article,
            'time_ago': time_ago
        })
    
    return render(request, 'app/indexed_articles.html', {
        'articles_with_time': articles_with_time,
        'articles': articles,
        'subjects': subjects_list,
        'journals': journals_list,
        'years': years_list,
        'page_obj': page_obj,
        'start_record': start_record,
        'end_record': end_record,
        'total_records': paginator.count,
        'per_page': per_page,
    })

def article_detail(request, article_id):
    """Article detail page view"""
    from django.shortcuts import get_object_or_404
    from datetime import date
    
    # Try to get article from database, otherwise use dummy data
    try:
        article = Article.objects.get(id=article_id)
        authors = article.authors.all().order_by('order')
        keywords_list = []
        if article.keywords:
            keywords_list = [k.strip() for k in article.keywords.split(',') if k.strip()]
        related_articles = Article.objects.filter(
            discipline=article.discipline
        ).exclude(id=article.id).order_by('-created_at')[:4]
        # Add image attribute to real articles (default image)
        for rel_article in related_articles:
            if not hasattr(rel_article, 'image'):
                rel_article.image = 'https://images.unsplash.com/photo-1451187580459-43490279c0fa?w=400&h=300&fit=crop'
        is_dummy = False
    except Article.DoesNotExist:
        # Create dummy article data based on article_id
        dummy_articles = {
            1: {
                'title': 'Blood-Pressure Targets in Comatose Survivors of Cardiac Arrest',
                'article_number': 'SIS123456AI141125',
                'discipline': 'Cardiology',
                'article_type': 'research',
                'abstract': 'This double-blind, randomized trial (BOX) compared mean arterial pressure targets of 77 mm Hg versus 63 mm Hg in comatose survivors of out-of-hospital cardiac arrest. The primary outcome was a composite of death from any cause or hospital discharge with a Cerebral Performance Category (CPC) score of 3 or 4 (indicating severe neurological disability) at 90 days. Secondary outcomes included neurological function, quality of life, and serious adverse events. A total of 789 patients were randomly assigned to the higher-target group (77 mm Hg) and 789 to the lower-target group (63 mm Hg). The primary outcome occurred in 48% of patients in the higher-target group and 52% in the lower-target group (relative risk, 0.92; 95% confidence interval [CI], 0.82 to 1.03; P=0.15). There were no significant differences in secondary outcomes. Serious adverse events occurred in 31% of patients in the higher-target group and 30% in the lower-target group. In comatose survivors of out-of-hospital cardiac arrest, targeting a mean arterial pressure of 77 mm Hg, as compared with 63 mm Hg, did not result in a significant difference in the composite outcome of death or severe neurological disability at 90 days.',
                'keywords': 'Blood pressure, cardiac arrest, comatose survivors, mean arterial pressure, neurological outcomes, randomized trial, intensive care',
                'doi': '10.1056/NEJMoa1900500',
                'publication_date': date(2024, 11, 14),
                'language': 'English',
                'volume': '10',
                'issue': '4',
                'pages': '123-130',
                'issn': '1234-5678',
                'views': 150,
                'citations': 5,
                'image': 'https://images.unsplash.com/photo-1551601651-2a8555f1a136?w=400&h=300&fit=crop',
                'authors': [
                    {'name': 'Dr. Sarah Johnson', 'affiliation': 'Department of Cardiology, University Medical Center', 'is_corresponding': True},
                    {'name': 'Dr. Michael Chen', 'affiliation': 'Cardiac Research Institute', 'is_corresponding': False},
                    {'name': 'Dr. Emily Rodriguez', 'affiliation': 'Emergency Medicine Department, City Hospital', 'is_corresponding': False},
                ]
            },
            2: {
                'title': 'Perceptions of School Administrators and Teachers on Educational Technology Integration',
                'article_number': 'SIS789012ED150126',
                'discipline': 'Education',
                'article_type': 'research',
                'abstract': 'The goal of the study is to review the perceptions of school administrators and teachers regarding the integration of educational technology in K-12 classrooms. This mixed-methods research examined factors influencing technology adoption, barriers to implementation, and the impact of professional development on technology integration practices. Data were collected through surveys, interviews, and classroom observations across 25 schools. Findings indicate that while administrators generally support technology integration, teachers face significant challenges including lack of training, insufficient technical support, and concerns about student distraction. The study recommends comprehensive professional development programs and increased technical infrastructure to support effective technology integration.',
                'keywords': 'Educational technology, technology integration, school administrators, teacher perceptions, professional development, K-12 education',
                'doi': '10.1016/j.edutech.2024.12345',
                'publication_date': date(2024, 12, 15),
                'language': 'English',
                'volume': '8',
                'issue': '3',
                'pages': '45-62',
                'issn': '2345-6789',
                'views': 89,
                'citations': 3,
                'image': 'https://images.unsplash.com/photo-1521737604893-d14cc237f11d?w=400&h=300&fit=crop',
                'authors': [
                    {'name': 'Dr. Robert Williams', 'affiliation': 'College of Education, State University', 'is_corresponding': True},
                    {'name': 'Dr. Lisa Anderson', 'affiliation': 'Educational Technology Research Center', 'is_corresponding': False},
                ]
            },
            3: {
                'title': 'Liberal or Restrictive Transfusion Strategy in Patients with Acute Myocardial Infarction',
                'article_number': 'SIS345678CA160127',
                'discipline': 'Cardiology',
                'article_type': 'research',
                'abstract': 'The SAHARA trial investigated the effect of a liberal versus restrictive red-cell transfusion strategy on clinical outcomes in patients with acute myocardial infarction and anemia. A total of 3,500 patients were randomly assigned to receive transfusions at a hemoglobin threshold of 10 g per deciliter (liberal strategy) or 8 g per deciliter (restrictive strategy). The primary outcome was a composite of death from any cause, recurrent myocardial infarction, or stroke at 30 days. Secondary outcomes included quality of life, functional status, and healthcare resource utilization. Results showed no significant difference in the primary composite outcome between groups (hazard ratio, 0.95; 95% CI, 0.82 to 1.10; P=0.48). However, the liberal strategy was associated with higher rates of transfusion-related complications. These findings suggest that a restrictive transfusion strategy may be appropriate for most patients with acute myocardial infarction and anemia.',
                'keywords': 'Transfusion strategy, myocardial infarction, anemia, red blood cells, clinical trial, cardiology',
                'doi': '10.1161/CIRCULATIONAHA.2024.98765',
                'publication_date': date(2024, 1, 16),
                'language': 'English',
                'volume': '12',
                'issue': '2',
                'pages': '234-251',
                'issn': '3456-7890',
                'views': 234,
                'citations': 12,
                'image': 'https://images.unsplash.com/photo-1576091160399-112ba8d25d1f?w=400&h=300&fit=crop',
                'authors': [
                    {'name': 'Dr. James Thompson', 'affiliation': 'Cardiovascular Research Institute', 'is_corresponding': True},
                    {'name': 'Dr. Patricia Martinez', 'affiliation': 'Department of Hematology, Medical Center', 'is_corresponding': False},
                    {'name': 'Dr. David Kim', 'affiliation': 'Clinical Trials Unit, University Hospital', 'is_corresponding': False},
                ]
            },
            4: {
                'title': 'Weekly Icodec versus Daily Glargine U100 in Type 2 Diabetes',
                'article_number': 'SIS456789EN170128',
                'discipline': 'Endocrinology',
                'article_type': 'research',
                'abstract': 'This phase 3a trial compared the efficacy and safety of once-weekly insulin icodec with once-daily insulin glargine U100 in adults with type 2 diabetes inadequately controlled on basal insulin. The study enrolled 1,200 participants who were randomly assigned to receive either icodec (700 units/mL) once weekly or glargine U100 (100 units/mL) once daily. The primary endpoint was the change in glycated hemoglobin (HbA1c) from baseline to week 26. Secondary endpoints included time in target glucose range, hypoglycemic events, and patient-reported outcomes. Results demonstrated non-inferiority of icodec compared to glargine U100, with mean HbA1c reductions of 1.2% and 1.1%, respectively. The incidence of hypoglycemia was similar between groups. Patient satisfaction scores were significantly higher in the icodec group, likely due to reduced injection frequency. These findings support once-weekly icodec as a convenient alternative to daily basal insulin therapy.',
                'keywords': 'Type 2 diabetes, insulin icodec, insulin glargine, weekly insulin, glycemic control, diabetes management',
                'doi': '10.2337/dc24-12345',
                'publication_date': date(2024, 1, 17),
                'language': 'English',
                'volume': '15',
                'issue': '1',
                'pages': '78-95',
                'issn': '4567-8901',
                'views': 312,
                'citations': 18,
                'image': 'https://images.unsplash.com/photo-1559757148-5c350d0d3c56?w=400&h=300&fit=crop',
                'authors': [
                    {'name': 'Dr. Jennifer Lee', 'affiliation': 'Diabetes Research Center', 'is_corresponding': True},
                    {'name': 'Dr. Christopher Brown', 'affiliation': 'Endocrinology Department, Regional Hospital', 'is_corresponding': False},
                ]
            },
        }
        
        dummy_data = dummy_articles.get(article_id, dummy_articles[1])
        
        # Create a simple object-like structure for template
        class DummyArticle:
            def __init__(self, data, art_id=None):
                self.id = art_id if art_id else article_id
                self.title = data['title']
                self.article_number = data.get('article_number', '')
                self.discipline = data['discipline']
                self.article_type = data['article_type']
                self.abstract = data['abstract']
                self.keywords = data['keywords']
                self.doi = data.get('doi', '')
                self.publication_date = data.get('publication_date')
                self.language = data.get('language', 'English')
                self.volume = data.get('volume', '')
                self.issue = data.get('issue', '')
                self.pages = data.get('pages', '')
                self.issn = data.get('issn', '')
                self.views = data.get('views', 0)
                self.citations = data.get('citations', 0)
                self.article_file = None
                self.image = data.get('image', '')
                self.created_at = data.get('publication_date', date.today())
            
            def get_article_type_display(self):
                type_map = {
                    'research': 'Research Article',
                    'review': 'Review Article',
                    'case_study': 'Case Study',
                    'short_communication': 'Short Communication',
                    'letter': 'Letter to Editor',
                }
                return type_map.get(self.article_type, 'Research Article')
        
        class DummyAuthor:
            def __init__(self, data):
                self.name = data['name']
                self.affiliation = data.get('affiliation', '')
                self.is_corresponding = data.get('is_corresponding', False)
        
        article = DummyArticle(dummy_data)
        authors = [DummyAuthor(auth) for auth in dummy_data['authors']]
        keywords_list = [k.strip() for k in dummy_data['keywords'].split(',') if k.strip()]
        
        # Dummy related articles
        related_dummy = []
        for rid in [2, 3, 4, 1]:
            if rid != article_id and rid in dummy_articles:
                rel_data = dummy_articles[rid]
                rel_article = DummyArticle(rel_data, art_id=rid)
                related_dummy.append(rel_article)
        
        related_articles = related_dummy[:4]
        is_dummy = True
    
    return render(request, 'app/article_detail.html', {
        'article': article,
        'authors': authors,
        'related_articles': related_articles,
        'keywords_list': keywords_list,
        'is_dummy': is_dummy,
    })

def article_certificate(request, article_id):
    """Generate and download article indexing certificate PDF using existing cert.pdf template"""
    from django.shortcuts import get_object_or_404
    from django.http import HttpResponse
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from datetime import datetime
    from PyPDF2 import PdfReader, PdfWriter
    from reportlab.pdfgen import canvas
    from io import BytesIO
    import os
    from django.conf import settings
    
    # Get article data (similar to article_detail)
    try:
        article = Article.objects.get(id=article_id)
        authors_list = article.authors.all().order_by('order')
        authors_names = ', '.join([author.name for author in authors_list]) if authors_list.exists() else article.authors_names or 'N/A'
        article_number = f"SIS{article.id:06d}AI{article.created_at.strftime('%d%m%y') if article.created_at else '000000'}"
        recorded_by = "Metascholar Limited"
        volume = article.volume or 'N/A'
        issue = article.issue or 'N/A'
        pages = article.pages or 'N/A'
        date_of_indexing = article.created_at.strftime('%B %d %Y') if article.created_at else datetime.now().strftime('%B %d %Y')
        article_link = request.build_absolute_uri(f'/indexed_articles/view/{article.id}/')
        article_title = article.title
    except Article.DoesNotExist:
        # Handle dummy articles
        from datetime import date
        dummy_articles = {
            1: {'title': 'Blood-Pressure Targets in Comatose Survivors of Cardiac Arrest', 'article_number': 'SIS123456AI141125', 'volume': '10', 'issue': '4', 'pages': '123-130', 'authors': [{'name': 'Dr. Sarah Johnson'}, {'name': 'Dr. Michael Chen'}, {'name': 'Dr. Emily Rodriguez'}]},
            2: {'title': 'Perceptions of School Administrators and Teachers on Educational Technology Integration', 'article_number': 'SIS789012ED150126', 'volume': '8', 'issue': '3', 'pages': '45-62', 'authors': [{'name': 'Dr. Robert Williams'}, {'name': 'Dr. Lisa Anderson'}]},
            3: {'title': 'Liberal or Restrictive Transfusion Strategy in Patients with Acute Myocardial Infarction', 'article_number': 'SIS345678CA160127', 'volume': '12', 'issue': '2', 'pages': '234-251', 'authors': [{'name': 'Dr. James Thompson'}, {'name': 'Dr. Patricia Martinez'}, {'name': 'Dr. David Kim'}]},
            4: {'title': 'Weekly Icodec versus Daily Glargine U100 in Type 2 Diabetes', 'article_number': 'SIS456789EN170128', 'volume': '15', 'issue': '1', 'pages': '78-95', 'authors': [{'name': 'Dr. Jennifer Lee'}, {'name': 'Dr. Christopher Brown'}]},
        }
        dummy_data = dummy_articles.get(article_id, dummy_articles[1])
        authors_names = ', '.join([a['name'] for a in dummy_data['authors']])
        article_number = dummy_data['article_number']
        recorded_by = "Metascholar Limited"
        volume = dummy_data['volume']
        issue = dummy_data['issue']
        pages = dummy_data['pages']
        date_of_indexing = datetime.now().strftime('%B %d %Y')
        article_link = request.build_absolute_uri(f'/indexed_articles/view/{article_id}/')
        article_title = dummy_data['title']
    
    # Load existing cert.pdf template
    cert_template_path = os.path.join(settings.BASE_DIR, 'app', 'static', 'app', 'docs', 'cert.pdf')
    
    # Create a BytesIO buffer for the overlay
    overlay_buffer = BytesIO()
    overlay_canvas = canvas.Canvas(overlay_buffer, pagesize=letter)
    width, height = letter
    
    # Set font and color for text overlay - use serif fonts for professional look
    overlay_canvas.setFillColor(HexColor('#000000'))
    
    # Use Times-Roman (serif) for a more professional certificate look
    from reportlab.lib.utils import simpleSplit
    
    # Starting Y position for article details section - moved up to avoid signatures
    start_y = height - 3.5 * inch
    line_height = 22  # Reduced line height for tighter spacing
    
    # Authors field with label - bold and stylish, very tight spacing like attached
    overlay_canvas.setFont("Times-Bold", 12)
    authors_label_x = 1.5 * inch
    overlay_canvas.drawString(authors_label_x, start_y, "Authors:")
    overlay_canvas.setFont("Times-Roman", 12)
    overlay_canvas.drawString(authors_label_x + 0.85 * inch, start_y, authors_names)
    
    # Article Number with label
    overlay_canvas.setFont("Times-Bold", 12)
    article_num_label_x = 1.5 * inch
    overlay_canvas.drawString(article_num_label_x, start_y - line_height, "Article Number:")
    overlay_canvas.setFont("Times-Roman", 12)
    overlay_canvas.drawString(article_num_label_x + 1.4 * inch, start_y - line_height, article_number)
    
    # Recorded By with label
    overlay_canvas.setFont("Times-Bold", 12)
    recorded_label_x = 1.5 * inch
    overlay_canvas.drawString(recorded_label_x, start_y - (line_height * 2), "Recorded By:")
    overlay_canvas.setFont("Times-Roman", 12)
    overlay_canvas.drawString(recorded_label_x + 1.2 * inch, start_y - (line_height * 2), recorded_by)
    
    # URL with label
    overlay_canvas.setFont("Times-Bold", 12)
    url_label_x = 1.5 * inch
    overlay_canvas.drawString(url_label_x, start_y - (line_height * 3), "URL:")
    overlay_canvas.setFont("Times-Roman", 12)
    overlay_canvas.drawString(url_label_x + 0.45 * inch, start_y - (line_height * 3), "www.scholarindexing.com")
    
    # Volume, Issue, Pages on same line with labels - very tight spacing
    vol_issue_y = start_y - (line_height * 4)
    overlay_canvas.setFont("Times-Bold", 12)
    vol_label_x = 1.5 * inch
    overlay_canvas.drawString(vol_label_x, vol_issue_y, "Volume:")
    overlay_canvas.setFont("Times-Roman", 12)
    overlay_canvas.drawString(vol_label_x + 0.7 * inch, vol_issue_y, volume)
    
    overlay_canvas.setFont("Times-Bold", 12)
    issue_label_x = 2.9 * inch
    overlay_canvas.drawString(issue_label_x, vol_issue_y, "Issue:")
    overlay_canvas.setFont("Times-Roman", 12)
    overlay_canvas.drawString(issue_label_x + 0.55 * inch, vol_issue_y, issue)
    
    overlay_canvas.setFont("Times-Bold", 12)
    pages_label_x = 4.7 * inch
    overlay_canvas.drawString(pages_label_x, vol_issue_y, "Pages:")
    overlay_canvas.setFont("Times-Roman", 12)
    overlay_canvas.drawString(pages_label_x + 0.6 * inch, vol_issue_y, pages)
    
    # Article Title Acknowledgment (wrapped text) - moved up, bold and larger font, with space before
    ack_y = start_y - (line_height * 5.2) - 20  # Added 20 points of space before
    ack_text = f'This certificate acknowledges that the article titled "{article_title}" has been successfully indexed and recorded by Metascholar Limited.'
    wrapped_ack = simpleSplit(ack_text, "Times-Bold", 13, width - 3*inch)
    for i, line in enumerate(wrapped_ack):
        overlay_canvas.setFont("Times-Bold", 13)
        overlay_canvas.drawString(1.5 * inch, ack_y - (i * 16), line)
    
    # Date of Indexing with label - moved up (adjusted for new line spacing), very tight spacing like attached
    date_y = ack_y - len(wrapped_ack) * 16 - 25
    overlay_canvas.setFont("Times-Bold", 12)
    date_label_x = 1.5 * inch
    overlay_canvas.drawString(date_label_x, date_y, "Date of Indexing:")
    overlay_canvas.setFont("Times-Roman", 12)
    overlay_canvas.drawString(date_label_x + 1.25 * inch, date_y, date_of_indexing)
    
    # Article Link with label - moved up, extremely tight spacing like attached
    link_y = date_y - line_height
    overlay_canvas.setFont("Times-Bold", 12)
    link_label_x = 1.5 * inch
    overlay_canvas.drawString(link_label_x, link_y, "Article Link:")
    overlay_canvas.setFont("Times-Roman", 11)
    link_wrapped = simpleSplit(article_link, "Times-Roman", 11, width - (link_label_x + 1.25*inch))
    for i, line in enumerate(link_wrapped):
        overlay_canvas.drawString(link_label_x + 1.25 * inch, link_y - (i * 13), line)
    
    overlay_canvas.save()
    
    # Read the existing cert.pdf template
    try:
        with open(cert_template_path, 'rb') as template_file:
            template_reader = PdfReader(template_file)
            template_writer = PdfWriter()
            
            # Get the first page of the template
            template_page = template_reader.pages[0]
            
            # Read the overlay PDF
            overlay_buffer.seek(0)
            overlay_reader = PdfReader(overlay_buffer)
            overlay_page = overlay_reader.pages[0]
            
            # Merge overlay onto template page
            template_page.merge_page(overlay_page)
            
            # Add the merged page to writer
            template_writer.add_page(template_page)
            
            # Create response
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="article_certificate_{article_id}.pdf"'
            
            # Write to response
            output_buffer = BytesIO()
            template_writer.write(output_buffer)
            output_buffer.seek(0)
            response.write(output_buffer.read())
            
            return response
    except FileNotFoundError:
        # If cert.pdf doesn't exist, create a new PDF (fallback)
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="article_certificate_{article_id}.pdf"'
        
        p = canvas.Canvas(response, pagesize=letter)
        width, height = letter
        
        # Draw basic certificate structure
        p.setFont("Helvetica-Bold", 20)
        p.drawCentredString(width/2, height - 1.5 * inch, "ARTICLE INDEXING CERTIFICATE")
        
        # Add article data
        p.setFont("Helvetica", 11)
        p.drawString(1.5 * inch, height - 3 * inch, f"Authors: {authors_names}")
        p.drawString(1.5 * inch, height - 3.25 * inch, f"Article Number: {article_number}")
        p.drawString(1.5 * inch, height - 3.5 * inch, f"Volume: {volume}, Issue: {issue}, Pages: {pages}")
        p.drawString(1.5 * inch, height - 4 * inch, f"Date of Indexing: {date_of_indexing}")
        p.drawString(1.5 * inch, height - 4.25 * inch, f"Article Link: {article_link}")
        
        # Article title acknowledgment
        ack_text = f'This certificate acknowledges that the article titled "{article_title}" has been successfully indexed and recorded by Metascholar Limited.'
        wrapped_ack = simpleSplit(ack_text, "Helvetica", 11, width - 3*inch)
        for i, line in enumerate(wrapped_ack):
            p.drawString(1.5 * inch, height - 4.5 * inch - (i * 15), line)
        
        p.showPage()
        p.save()
        
        return response

@check_page_enabled('enable_indexed_journals_page')
def indexed_journals(request):
    """Indexed Journals page view"""
    journals = Journal.objects.all().order_by('-created_at')
    return render(request, 'app/indexed_journals.html', {'journals': journals})

def journal_detail(request, journal_id):
    """Journal detail page view"""
    from django.shortcuts import get_object_or_404
    
    journal = get_object_or_404(Journal, id=journal_id)
    
    # Get chief editor (first editor or first by order)
    chief_editor = journal.editors.first()
    
    # Get similar journals (same subject area, excluding current)
    similar_journals = Journal.objects.filter(
        subject_area=journal.subject_area
    ).exclude(id=journal.id).order_by('-created_at')[:4]
    
    # If not enough similar journals, get any other journals
    if similar_journals.count() < 4:
        other_journals = Journal.objects.exclude(id=journal.id).exclude(
            id__in=[j.id for j in similar_journals]
        ).order_by('-created_at')[:4 - similar_journals.count()]
        similar_journals = list(similar_journals) + list(other_journals)
    
    # Generate journal number (SIS + ID + JI + timestamp)
    if journal.created_at:
        date_str = journal.created_at.strftime('%y%m%d')
    else:
        date_str = '000000'
    journal_number = f"SIS{journal.id:06d}.JI{date_str}"
    
    return render(request, 'app/journal_detail.html', {
        'journal': journal,
        'chief_editor': chief_editor,
        'similar_journals': similar_journals,
        'journal_number': journal_number,
    })

@check_page_enabled('enable_project_archive_page')
def project_archive(request):
    """Project | Research Archive page view"""
    projects = Project.objects.all()
    
    # Get filter parameters from GET request
    search_query = request.GET.get('search', '').strip()
    category = request.GET.get('category', '').strip()
    institution = request.GET.get('institution', '').strip()
    status = request.GET.get('status', '').strip()
    year = request.GET.get('year', '').strip()
    sort_by = request.GET.get('sort', 'recent')
    
    # Apply search filter
    if search_query:
        projects = projects.filter(
            Q(project_title__icontains=search_query) |
            Q(description__icontains=search_query) |
            Q(institution__icontains=search_query) |
            Q(category__icontains=search_query)
        )
    
    # Apply category filter
    if category:
        projects = projects.filter(category__icontains=category)
    
    # Apply institution filter
    if institution:
        projects = projects.filter(institution__icontains=institution)
    
    # Apply status filter
    if status:
        projects = projects.filter(status=status)
    
    # Apply year filter
    if year:
        projects = projects.filter(created_at__year=year)
    
    # Apply sorting
    if sort_by == 'recent':
        projects = projects.order_by('-created_at')
    elif sort_by == 'oldest':
        projects = projects.order_by('created_at')
    elif sort_by == 'title':
        projects = projects.order_by('project_title')
    else:
        projects = projects.order_by('-created_at')
    
    # Get unique values for filter dropdowns
    institutions = Project.objects.values_list('institution', flat=True).distinct().order_by('institution')
    categories = Project.objects.values_list('category', flat=True).distinct().order_by('category')
    years = Project.objects.dates('created_at', 'year', order='DESC')
    
    return render(request, 'app/project_archive.html', {
        'projects': projects,
        'institutions': institutions,
        'categories': categories,
        'years': years,
        'current_search': search_query,
        'current_category': category,
        'current_institution': institution,
        'current_status': status,
        'current_year': year,
        'current_sort': sort_by,
    })

@check_page_enabled('enable_project_archive_page')
def project_detail(request, project_id):
    """Project detail page view"""
    project = get_object_or_404(Project, id=project_id)
    
    # Increment views
    project.views += 1
    project.save(update_fields=['views'])
    
    # Get project author (submitted_by)
    author = project.submitted_by
    
    # Parse additional_info to extract price, chapters, pages
    price = project.price_usd
    chapters = "1-3"
    pages = "1-90"
    
    if project.additional_info:
        for line in project.additional_info.split('\n'):
            if 'Price in USD:' in line:
                try:
                    price = float(line.split(':')[1].strip())
                except:
                    pass
            elif 'Chapters:' in line:
                chapters = line.split(':')[1].strip()
            elif 'Pages:' in line:
                pages = line.split(':')[1].strip()
    
    # Get similar projects (same category, excluding current)
    similar_projects = Project.objects.filter(
        category=project.category
    ).exclude(id=project.id).order_by('-created_at')[:3]
    
    # If not enough similar projects, get any other projects
    if similar_projects.count() < 3:
        other_projects = Project.objects.exclude(id=project.id).exclude(
            id__in=[p.id for p in similar_projects]
        ).order_by('-created_at')[:3 - similar_projects.count()]
        similar_projects = list(similar_projects) + list(other_projects)
    
    # Parse chapters and pages for similar projects
    similar_projects_data = []
    for sp in similar_projects:
        sp_chapters = "1-3"
        sp_pages = "1-90"
        if sp.additional_info:
            for line in sp.additional_info.split('\n'):
                if 'Chapters:' in line:
                    sp_chapters = line.split(':')[1].strip()
                elif 'Pages:' in line:
                    sp_pages = line.split(':')[1].strip()
        similar_projects_data.append({
            'project': sp,
            'chapters': sp_chapters,
            'pages': sp_pages,
        })
    
    # Check if user has already paid (via email in session or payment record)
    has_paid = False
    payment_email = request.session.get('payment_email', '')
    if payment_email:
        has_paid = ProjectPayment.objects.filter(
            project=project,
            email=payment_email,
            payment_status='completed'
        ).exists()
    
    # Get Paystack public key - check environment variables first, then settings.py
    import os
    from django.conf import settings as django_settings
    
    # Check environment variables (VITE_ prefix or direct)
    paystack_public_key = (
        os.getenv('VITE_PAYSTACK_PUBLIC_KEY') or 
        os.getenv('PAYSTACK_PUBLIC_KEY') or 
        getattr(django_settings, 'PAYSTACK_PUBLIC_KEY', None)
    )
    
    # Validate the key format
    if not paystack_public_key or not isinstance(paystack_public_key, str) or not paystack_public_key.startswith('pk_'):
        # Fallback to the known test key
        paystack_public_key = 'pk_test_af37d26c0fa360522c4e66495f3877e498c18850'
    
    return render(request, 'app/project_detail.html', {
        'project': project,
        'author': author,
        'price': price,
        'chapters': chapters,
        'pages': pages,
        'similar_projects': similar_projects_data,
        'has_paid': has_paid,
        'payment_email': payment_email,
        'paystack_public_key': paystack_public_key,
    })

def verify_project_payment(request, project_id):
    """Verify payment and send document via email"""
    import requests
    import json
    from django.http import JsonResponse
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            email = data.get('email', '').strip()
            payment_ref = data.get('reference', '').strip()
            
            if not email or not payment_ref:
                return JsonResponse({'error': 'Email and payment reference are required'}, status=400)
            
            project = get_object_or_404(Project, id=project_id)
            
            # Verify payment with Paystack - check environment variables first, then settings.py
            from django.conf import settings as django_settings
            
            # Check environment variables (VITE_ prefix or direct)
            paystack_secret = (
                os.getenv('VITE_PAYSTACK_SECRET_KEY') or 
                os.getenv('PAYSTACK_SECRET_KEY') or 
                getattr(django_settings, 'PAYSTACK_SECRET_KEY', None)
            )
            
            # Fallback to default if not found
            if not paystack_secret:
                paystack_secret = 'sk_test_185fc53d96addab7232060c86f4221918ab59d1c'
            
            url = f'https://api.paystack.co/transaction/verify/{payment_ref}'
            headers = {
                'Authorization': f'Bearer {paystack_secret}',
            }
            
            response = requests.get(url, headers=headers)
            result = response.json()
            
            if result.get('status') and result.get('data', {}).get('status') == 'success':
                # Payment successful
                amount = float(result['data']['amount']) / 100  # Convert from pesewas to GHS, then to USD
                amount_usd = amount / 13.5  # Approximate conversion
                
                # Create or update payment record
                payment, created = ProjectPayment.objects.get_or_create(
                    payment_reference=payment_ref,
                    defaults={
                        'project': project,
                        'email': email,
                        'amount_paid': amount_usd,
                        'payment_status': 'completed',
                    }
                )
                
                if not created:
                    payment.payment_status = 'completed'
                    payment.save()
                
                # Send document via email if not already sent
                if not payment.document_sent and project.project_file:
                    try:
                        from django.core.mail import EmailMessage
                        from django.template.loader import render_to_string
                        from django.conf import settings as django_settings
                        import traceback
                        
                        subject = f'Your Project Document - {project.project_title}'
                        html_message = render_to_string('app/emails/project_document.html', {
                            'project': project,
                            'email': email,
                        })
                        
                        # Get email settings
                        from_email = getattr(django_settings, 'DEFAULT_FROM_EMAIL', 'noreply@scholarindex.com')
                        
                        email_msg = EmailMessage(
                            subject,
                            html_message,
                            from_email,
                            [email],
                        )
                        email_msg.content_subtype = 'html'
                        
                        # Attach file - handle different storage backends
                        try:
                            # Try to get file path (works for local storage)
                            if hasattr(project.project_file, 'path'):
                                file_path = project.project_file.path
                                if os.path.exists(file_path):
                                    email_msg.attach_file(file_path)
                                else:
                                    raise Exception(f"File not found at path: {file_path}")
                            else:
                                # For cloud storage (S3, Cloudinary, etc.), read the file
                                project.project_file.open('rb')
                                file_content = project.project_file.read()
                                file_name = project.project_file.name.split('/')[-1]
                                
                                # Determine content type based on file extension
                                import mimetypes
                                content_type, _ = mimetypes.guess_type(file_name)
                                if not content_type:
                                    if file_name.endswith('.pdf'):
                                        content_type = 'application/pdf'
                                    elif file_name.endswith(('.doc', '.docx')):
                                        content_type = 'application/msword'
                                    elif file_name.endswith('.zip'):
                                        content_type = 'application/zip'
                                    else:
                                        content_type = 'application/octet-stream'
                                
                                email_msg.attach(file_name, file_content, content_type)
                                project.project_file.close()
                        except Exception as file_error:
                            raise Exception(f"Failed to attach file: {str(file_error)}")
                        
                        # Send email
                        email_sent = email_msg.send(fail_silently=False)
                        
                        if email_sent:
                            payment.document_sent = True
                            payment.save()
                            
                            # Increment downloads
                            project.downloads += 1
                            project.save(update_fields=['downloads'])
                            
                            # Store email in session
                            request.session['payment_email'] = email
                        else:
                            raise Exception("Email send() returned 0")
                        
                    except Exception as e:
                        import logging
                        logger = logging.getLogger(__name__)
                        logger.error(f"Error sending email: {str(e)}")
                        logger.error(traceback.format_exc())
                        return JsonResponse({
                            'error': f'Payment verified but email failed to send: {str(e)}'
                        }, status=500)
                
                return JsonResponse({
                    'success': True,
                    'message': 'Payment verified successfully. Document sent to your email.',
                })
            else:
                return JsonResponse({'error': 'Payment verification failed'}, status=400)
                
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Invalid request method'}, status=405)

@login_required
def upload_article(request):
    """Upload Article page view"""
    if request.method == 'POST':
        try:
            # Get form data
            title = request.POST.get('title', '').strip()
            discipline = request.POST.get('discipline', '').strip()
            abstract = request.POST.get('abstract', '').strip()
            keywords = request.POST.get('keywords', '').strip()
            authors_names = request.POST.get('authors_names', '').strip()
            year_of_publication = request.POST.get('year_of_publication', '').strip()
            volume = request.POST.get('volume', '').strip()
            issue = request.POST.get('issue', '').strip()
            pages = request.POST.get('pages', '').strip()
            journal_name = request.POST.get('journal_name', '').strip()
            country_of_publication = request.POST.get('country_of_publication', '').strip()
            issn_or_doi = request.POST.get('issn_or_doi', '').strip()
            cover_image = request.FILES.get('cover_image')
            article_file = request.FILES.get('article_file')
            
            # Validate required fields
            if not title:
                messages.error(request, 'Article title is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not discipline:
                messages.error(request, 'Subject is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not abstract:
                messages.error(request, 'Abstract is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not keywords:
                messages.error(request, 'Keywords are required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not authors_names:
                messages.error(request, 'Authors names are required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not year_of_publication:
                messages.error(request, 'Year of publication is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not volume:
                messages.error(request, 'Volume is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not issue:
                messages.error(request, 'Issue is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not pages:
                messages.error(request, 'Page numbers are required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not journal_name:
                messages.error(request, 'Journal name is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not country_of_publication:
                messages.error(request, 'Country of publication is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not issn_or_doi:
                messages.error(request, 'ISSN or DOI is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            if not article_file:
                messages.error(request, 'Article file is required.')
                return render(request, 'app/upload_article.html', {'form': ArticleForm()})
            
            # Set default article_type to 'research' if not provided
            article_type = 'research'
            
            # Create publication date from year
            from datetime import datetime
            try:
                year_int = int(year_of_publication)
                pub_date = datetime(year_int, 1, 1).date()
            except:
                pub_date = None
            
            # Determine if issn_or_doi is DOI or ISSN
            doi = ''
            issn = ''
            if issn_or_doi.startswith('10.') or 'doi' in issn_or_doi.lower():
                doi = issn_or_doi
            else:
                issn = issn_or_doi
            
            article = Article.objects.create(
                title=title,
                article_type=article_type,
                discipline=discipline,
                abstract=abstract,
                keywords=keywords,
                language='English',  # Default language
                publication_date=pub_date,
                doi=doi,
                article_file=article_file,
                cover_image=cover_image,
                volume=volume,
                issue=issue,
                pages=pages,
                journal_name=journal_name,
                country_of_publication=country_of_publication,
                year_of_publication=year_int if year_of_publication else None,
                authors_names=authors_names,
                submitted_by=request.user,
                status='pending'
            )
            
            # Handle authors from comma-separated names
            if authors_names:
                author_list = [name.strip() for name in authors_names.split(',') if name.strip()]
                for index, author_name in enumerate(author_list):
                    ArticleAuthor.objects.create(
                        article=article,
                        name=author_name,
                        email=request.user.email if index == 0 else '',  # Use user email for first author
                        affiliation='',
                        is_corresponding=(index == 0),
                        order=index
                    )
            
            messages.success(request, 'Article uploaded successfully! It will be reviewed before being published.')
            return redirect('app:indexed_articles')
            
        except Exception as e:
            messages.error(request, f'Error uploading article: {str(e)}')
            import traceback
            print(traceback.format_exc())
    else:
        form = ArticleForm()
    
    return render(request, 'app/upload_article.html', {'form': form})

@login_required
def register_journal(request):
    """Register Journal page view"""
    if request.method == 'POST':
        # Validate captcha (handled by JavaScript, but verify server-side too)
        captcha_input = request.POST.get('captcha_input', '').strip().upper()
        captcha_code = request.session.get('captcha_code', '')
        
        if not captcha_code:
            # Generate captcha if missing
            import random
            import string
            captcha_code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=5))
            request.session['captcha_code'] = captcha_code
        
        if captcha_input != captcha_code:
            messages.error(request, 'Invalid captcha code. Please try again.')
            # Regenerate captcha for next attempt
            import random
            import string
            new_captcha = ''.join(random.choices(string.ascii_uppercase + string.digits, k=5))
            request.session['captcha_code'] = new_captcha
            return render(request, 'app/register_journal.html', {
                'form': JournalForm(),
                'captcha_code': new_captcha,
            })
        
        # Create a mutable copy of POST data
        post_data = request.POST.copy()
        
        # Handle radio button values for open_access and peer_review
        open_access_value = post_data.get('open_access', 'false')
        peer_review_value = post_data.get('peer_review', 'false')
        post_data['open_access'] = True if open_access_value == 'true' else False
        post_data['peer_review'] = True if peer_review_value == 'true' else False
        
        form = JournalForm(post_data, request.FILES)
        if form.is_valid():
            journal = form.save(commit=False)
            journal.submitted_by = request.user
            # Ensure boolean fields are set correctly
            journal.open_access = open_access_value == 'true'
            journal.peer_review = peer_review_value == 'true'
            journal.save()
            
            # Handle Chief Editor
            chief_editor_name = request.POST.get('chief_editor_name', '').strip()
            chief_editor_email = request.POST.get('chief_editor_email', '').strip()
            
            if chief_editor_name and chief_editor_email:
                JournalEditor.objects.create(
                    journal=journal,
                    name=chief_editor_name,
                    email=chief_editor_email,
                    role='Chief Editor',
                    order=0
                )
            
            # Clear captcha after successful submission
            request.session.pop('captcha_code', None)
            
            messages.success(request, 'Journal registered successfully!')
            return redirect('app:indexed_journals')
        else:
            # Regenerate captcha if form validation fails
            import random
            import string
            new_captcha = ''.join(random.choices(string.ascii_uppercase + string.digits, k=5))
            request.session['captcha_code'] = new_captcha
            
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = JournalForm()
        # Generate initial captcha code if not exists
        if 'captcha_code' not in request.session:
            import random
            import string
            captcha_code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=5))
            request.session['captcha_code'] = captcha_code
    
    # Always get captcha code from session
    captcha_code = request.session.get('captcha_code', '')
    if not captcha_code:
        import random
        import string
        captcha_code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=5))
        request.session['captcha_code'] = captcha_code
    
    return render(request, 'app/register_journal.html', {
        'form': form,
        'captcha_code': captcha_code,
    })

@check_page_enabled('enable_directory_researchers_page')
def directory_researchers(request):
    """Directory of Researchers page view"""
    researchers = DirectoryApplication.objects.filter(terms_accepted=True)
    
    # Get filter parameters from GET request
    search_query = request.GET.get('search', '').strip()
    country = request.GET.get('country', '').strip()
    institution = request.GET.get('institution', '').strip()
    discipline = request.GET.get('discipline', '').strip()
    sort_by = request.GET.get('sort', 'recent')
    
    # Apply search filter
    if search_query:
        researchers = researchers.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(institution__icontains=search_query) |
            Q(position__icontains=search_query) |
            Q(research_areas__icontains=search_query)
        )
    
    # Apply country filter
    if country:
        researchers = researchers.filter(country__icontains=country)
    
    # Apply institution filter
    if institution:
        researchers = researchers.filter(institution__icontains=institution)
    
    # Apply discipline/research area filter
    if discipline:
        researchers = researchers.filter(research_areas__icontains=discipline)
    
    # Apply sorting
    if sort_by == 'name':
        researchers = researchers.order_by('first_name', 'last_name')
    elif sort_by == 'recent':
        researchers = researchers.order_by('-created_at')
    else:
        researchers = researchers.order_by('-created_at')
    
    # Get unique values for filter dropdowns
    from django.db.models import Count
    countries = DirectoryApplication.objects.filter(terms_accepted=True).exclude(country='').values('country').annotate(count=Count('id')).order_by('country')
    countries_list = [{'name': c['country'], 'count': c['count']} for c in countries if c['country']]
    
    institutions = DirectoryApplication.objects.filter(terms_accepted=True).exclude(institution='').values('institution').annotate(count=Count('id')).order_by('institution')
    institutions_list = [{'name': i['institution'], 'count': i['count']} for i in institutions if i['institution']]
    
    # Predefined areas of interest
    areas_of_interest = [
        'Music education, Community music, Applied ethnomusicology, Secondary school music curricula, Intercultural relations and organised cultural encounters',
        'Brain Computer, Interface Biomedical, Signal Processing, Machine Learning, Deep Learning, Artificial Intelligence',
        'Artificial Intelligence, Machine Learning, Deep Learning',
        'Populations of bacteria, Evolution of symbioses',
        'Water security, water resources management, climate change impact',
        'Smoking cessation, Harm reduction, Evidence synthesis',
        'Time Series Econometrics, Forecasting, Macroeconometrics',
        'Strategic Communications, Marketing & Communications',
        'Chartered Accountant',
        'Chemical ecology and management of agricultural pests for sustainable crop production, postharvest technology for food security',
        'Theology, Environmental health, Community health, Epidemiology',
        'Education, Curriculum studies, leadership studies, social policy, teaching',
        'Organisation and Human Development, Management, General Education',
        'Epidemiology, Occupational Health & Safety',
        'Data analysis, EHR global health, health policy, and information management research skills/knowledge',
        'Non-communicable diseases, mental health care',
        'Public Health',
        'Health Economics, Behavioral Economics, Public Economics, Industrial Organization',
        'Finance, Banking, Economics',
        'Time Series Analysis, Empirical Finance, Financial Markets, Corporate Finance, Digital Assets & FinTech',
    ]
    
    # Count researchers for each area of interest (show all areas, sorted by count descending)
    areas_with_count = []
    for area in areas_of_interest:
        count = DirectoryApplication.objects.filter(terms_accepted=True, research_areas__icontains=area).count()
        areas_with_count.append({'name': area, 'count': count})
    
    # Sort by count descending (highest to smallest)
    areas_with_count.sort(key=lambda x: x['count'], reverse=True)
    
    return render(request, 'app/directory_researchers.html', {
        'researchers': researchers,
        'countries': countries_list,
        'institutions': institutions_list,
        'areas_of_interest': areas_with_count,
        'current_search': search_query,
        'current_country': country,
        'current_institution': institution,
        'current_discipline': discipline,
        'current_sort': sort_by,
    })

@login_required
def upload_project(request):
    """Upload | Archive Project page view"""
    if request.method == 'POST':
        # Create a mutable copy of POST data
        post_data = request.POST.copy()
        
        # Map new form fields to model fields
        # subject -> store in additional_info
        # price_usd, chapters, pages -> store in additional_info
        additional_info_parts = []
        
        subject = request.POST.get('subject', '')
        price_usd = request.POST.get('price_usd', '')
        chapters = request.POST.get('chapters', '')
        pages = request.POST.get('pages', '')
        
        if subject:
            additional_info_parts.append(f"Subject: {subject}")
        if price_usd:
            additional_info_parts.append(f"Price in USD: {price_usd}")
        if chapters:
            additional_info_parts.append(f"Chapters: {chapters}")
        if pages:
            additional_info_parts.append(f"Pages: {pages}")
        
        # Set project_type and status to defaults if not provided
        if not post_data.get('project_type'):
            post_data['project_type'] = 'research'
        if not post_data.get('status'):
            post_data['status'] = 'active'
        if not post_data.get('institution'):
            post_data['institution'] = 'Not specified'
        
        # Set additional_info
        if additional_info_parts:
            post_data['additional_info'] = '\n'.join(additional_info_parts)
        
        form = ProjectForm(post_data, request.FILES)
        if form.is_valid():
            project = form.save(commit=False)
            project.submitted_by = request.user
            project.save()
            
            # Handle contributors
            contributor_count = int(request.POST.get('contributor_count', 0))
            for i in range(contributor_count):
                contributor_name = request.POST.get(f'contributor_{i}_name')
                contributor_email = request.POST.get(f'contributor_{i}_email')
                contributor_affiliation = request.POST.get(f'contributor_{i}_affiliation', '')
                contributor_role = request.POST.get(f'contributor_{i}_role', '')
                
                if contributor_name and contributor_email:
                    ProjectContributor.objects.create(
                        project=project,
                        name=contributor_name,
                        email=contributor_email,
                        affiliation=contributor_affiliation,
                        role=contributor_role,
                        order=i
                    )
            
            messages.success(request, 'Project uploaded and archived successfully!')
            return redirect('app:project_archive')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = ProjectForm()
    
    return render(request, 'app/upload_project.html', {'form': form})

@check_page_enabled('enable_council_members_page')
def council_members(request):
    """Council Members page view"""
    return render(request, 'app/council_members.html')

@check_page_enabled('enable_team_members_page')
def team_members(request):
    """Team Members page view"""
    return render(request, 'app/team_members.html')

@check_page_enabled('enable_donate_page')
def donate(request):
    """Donate page view"""
    from django.conf import settings
    return render(request, 'app/donate.html', {
        'paystack_public_key': settings.PAYSTACK_PUBLIC_KEY,
    })

def sponsors(request):
    """Sponsors page view - displays list of contributors/sponsors"""
    # Sponsor data with actual logo URLs
    sponsors_data = {
        'premier': [
            {
                'name': 'University of Energy and Natural Resources',
                'logo': 'https://res.cloudinary.com/dmqizfpyz/image/upload/v1763228952/uenr_kxmfdz.png',
            },
            {
                'name': 'U.S. Department of the Interior',
                'logo': 'https://res.cloudinary.com/dmqizfpyz/image/upload/v1763228952/premier_yploja.png',
            },
        ],
        'sustaining': [
            {
                'name': 'IEEE',
                'logo': 'https://res.cloudinary.com/dmqizfpyz/image/upload/v1763228951/ieee_idrslf.gif',
            },
            {
                'name': 'Catholic University of Ghana',
                'logo': 'https://res.cloudinary.com/dmqizfpyz/image/upload/v1763228951/cug_tsz1xw.png',
            },
            {
                'name': 'Sunyani Technical University',
                'logo': 'https://res.cloudinary.com/dmqizfpyz/image/upload/v1763228952/stu_uvpu4j.png',
            },
        ],
        'basic': [
            {
                'name': 'International Journal of Multidisciplinary Studies and Innovative Research',
                'logo': 'https://res.cloudinary.com/dmqizfpyz/image/upload/v1763228952/ijms_pdlmyd.jpg',
            },
        ],
    }
    
    return render(request, 'app/sponsors.html', {
        'sponsors': sponsors_data,
    })

def initialize_payment(request):
    """Initialize Paystack payment"""
    import requests
    from django.conf import settings
    from django.http import JsonResponse
    import json
    
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            amount_usd_cents = float(data.get('amount', 0))  # Amount in USD cents
            email = data.get('email', '')
            name = data.get('name', 'Donor')
            
            if amount_usd_cents <= 0:
                return JsonResponse({'error': 'Invalid amount'}, status=400)
            
            # Convert USD cents to GHS pesewas (1 USD = ~13.5 GHS)
            # First convert cents to dollars, then to GHS, then to pesewas
            amount_usd = amount_usd_cents / 100
            amount_ghs = amount_usd * 13.5  # Approximate conversion rate
            amount_pesewas = int(amount_ghs * 100)  # Convert to pesewas
            
            # Initialize Paystack transaction
            url = 'https://api.paystack.co/transaction/initialize'
            headers = {
                'Authorization': f'Bearer {settings.PAYSTACK_SECRET_KEY}',
                'Content-Type': 'application/json'
            }
            payload = {
                'email': email,
                'amount': amount_pesewas,
                'currency': 'GHS',
                'reference': f'sponsor_{request.user.id if request.user.is_authenticated else "guest"}_{int(time.time())}',
                'metadata': {
                    'name': name,
                    'custom_fields': [
                        {
                            'display_name': 'Donation Type',
                            'variable_name': 'donation_type',
                            'value': 'Sponsor'
                        }
                    ]
                }
            }
            
            response = requests.post(url, headers=headers, json=payload)
            result = response.json()
            
            if result.get('status'):
                return JsonResponse({
                    'authorization_url': result['data']['authorization_url'],
                    'access_code': result['data']['access_code'],
                    'reference': result['data']['reference']
                })
            else:
                return JsonResponse({'error': result.get('message', 'Payment initialization failed')}, status=400)
                
        except Exception as e:
            return JsonResponse({'error': str(e)}, status=500)
    
    return JsonResponse({'error': 'Invalid request method'}, status=405)

def request_membership(request):
    """Request For SIS Member page view"""
    from app.forms import MembershipRequestForm
    
    if request.method == 'POST':
        # Get form data and prepare for form
        full_name = request.POST.get('full_name', '').strip()
        email = request.POST.get('email', '').strip()
        country = request.POST.get('country', '').strip()
        institution = request.POST.get('institution', '').strip()
        aspiring_position = request.POST.get('aspiring_position', '').strip()
        about_yourself = request.POST.get('about_yourself', '').strip()
        terms = request.POST.get('terms') == 'on'
        profile_picture = request.FILES.get('profile_picture')
        
        # Validation
        if not full_name:
            messages.error(request, 'Please enter your first name and surname.')
        elif not email:
            messages.error(request, 'Please enter your email address.')
        elif not country:
            messages.error(request, 'Please select a country.')
        elif not institution:
            messages.error(request, 'Please enter your institution.')
        elif not aspiring_position:
            messages.error(request, 'Please select an aspiring position.')
        elif not about_yourself:
            messages.error(request, 'Please tell us more about yourself.')
        elif not terms:
            messages.error(request, 'Please agree to the Terms and Conditions and Privacy policy.')
        elif not profile_picture:
            messages.error(request, 'Please upload a profile picture.')
        else:
            # Split full name into first and last name
            name_parts = full_name.split(maxsplit=1)
            first_name = name_parts[0] if name_parts else ''
            last_name = name_parts[1] if len(name_parts) > 1 else ''
            
            # Prepare form data
            form_data = {
                'first_name': first_name,
                'last_name': last_name,
                'email': email,
                'phone': '',
                'country': country,
                'institution': institution,
                'position': aspiring_position,
                'membership_type': 'individual',
                'research_interests': about_yourself,
                'terms_accepted': terms,
            }
            
            form = MembershipRequestForm(form_data, request.FILES)
            if form.is_valid():
                membership = form.save(commit=False)
                membership.submitted_by = request.user if request.user.is_authenticated else None
                membership.save()
                messages.success(request, 'Your membership request has been submitted successfully!')
                return redirect('app:landing')
            else:
                for field, errors in form.errors.items():
                    for error in errors:
                        messages.error(request, f'{field}: {error}')
    
    return render(request, 'app/request_membership.html')

def apply_directory(request):
    """Request For SIS Member - Directory Application page view"""
    from app.forms import DirectoryApplicationForm
    
    if request.method == 'POST':
        # Get form data
        full_name = request.POST.get('full_name', '').strip()
        phone = request.POST.get('phone', '').strip()
        email = request.POST.get('email', '').strip()
        country = request.POST.get('country', '').strip()
        institution = request.POST.get('institution', '').strip()
        designation = request.POST.get('designation', '').strip()
        google_scholar = request.POST.get('google_scholar', '').strip()
        areas_of_interest = request.POST.get('areas_of_interest', '').strip()
        terms = request.POST.get('terms') == 'on'
        profile_picture = request.FILES.get('profile_picture')
        cv_file = request.FILES.get('cv_file')
        
        # Validation
        if not full_name:
            messages.error(request, 'Please enter your first name and surname.')
        elif not phone:
            messages.error(request, 'Please enter your phone number.')
        elif not email:
            messages.error(request, 'Please enter your email address.')
        elif not country:
            messages.error(request, 'Please enter your country.')
        elif not institution:
            messages.error(request, 'Please enter your institution.')
        elif not designation:
            messages.error(request, 'Please enter your designation.')
        elif not google_scholar:
            messages.error(request, 'Please enter your Google Scholar Profile Link.')
        elif not areas_of_interest:
            messages.error(request, 'Please enter your areas of interest.')
        elif not terms:
            messages.error(request, 'Please agree to the Terms and Conditions and Privacy policy.')
        elif not profile_picture:
            messages.error(request, 'Please upload a profile picture.')
        elif not cv_file:
            messages.error(request, 'Please upload your CV file.')
        else:
            # Split full name into first and last name
            name_parts = full_name.split(maxsplit=1)
            first_name = name_parts[0] if name_parts else ''
            last_name = name_parts[1] if len(name_parts) > 1 else ''
            
            # Prepare form data
            form_data = {
                'first_name': first_name,
                'last_name': last_name,
                'email': email,
                'phone': phone,
                'country': country,
                'institution': institution,
                'position': designation,
                'research_areas': areas_of_interest,
                'google_scholar_link': google_scholar,
                'terms_accepted': terms,
            }
            
            form = DirectoryApplicationForm(form_data, request.FILES)
            if form.is_valid():
                application = form.save(commit=False)
                application.submitted_by = request.user if request.user.is_authenticated else None
                application.save()
                messages.success(request, 'Your application has been submitted successfully!')
                return redirect('app:directory_researchers')
            else:
                for field, errors in form.errors.items():
                    for error in errors:
                        messages.error(request, f'{field}: {error}')
    
    return render(request, 'app/apply_directory.html')

@check_page_enabled('enable_hall_of_fame_page')
def hall_of_fame(request):
    """Hall of Fame listing page view - Eminent Personalities"""
    honorees = HallOfFameApplication.objects.filter(terms_accepted=True).order_by('-created_at')
    
    # Get unique years from created_at (year inducted)
    from django.db.models import Count
    from django.db.models.functions import ExtractYear
    years = HallOfFameApplication.objects.filter(terms_accepted=True).annotate(
        year=ExtractYear('created_at')
    ).values('year').annotate(count=Count('id')).order_by('-year')
    
    # Get unique disciplines with counts from database
    disciplines_with_count = HallOfFameApplication.objects.filter(terms_accepted=True).values('nominee_position').annotate(
        count=Count('id')
    ).order_by('-count', 'nominee_position')
    
    # Convert to list format for template
    disciplines_list = [{'name': item['nominee_position'], 'count': item['count']} for item in disciplines_with_count]
    
    # Get categories from application_type (for now, using application_type as category)
    # In production, you might want to add a separate category field to the model
    categories_with_count = HallOfFameApplication.objects.filter(terms_accepted=True).values('application_type').annotate(
        count=Count('id')
    ).order_by('-count', 'application_type')
    
    # Convert to list format for template
    categories_list = [{'name': item['application_type'].replace('_', ' ').title(), 'count': item['count']} for item in categories_with_count]
    
    # If no categories found, use default sample data
    if not categories_list:
        categories_list = [
            {'name': 'Eminent Individuals', 'count': 0},
            {'name': 'Celebrities', 'count': 0},
            {'name': 'Sport', 'count': 0},
            {'name': 'Inspirational people', 'count': 0},
            {'name': 'Famous People', 'count': 0},
            {'name': 'Social and Behavioral Sciences', 'count': 0},
            {'name': 'Famous Historical Figures', 'count': 0},
        ]
    
    return render(request, 'app/hall_of_fame.html', {
        'honorees': honorees,
        'years': years,
        'disciplines': disciplines_list,
        'categories': categories_list,
    })

def hall_of_fame_apply(request):
    """Apply For Hall of Fame page view"""
    if request.method == 'POST':
        # Handle full name splitting
        post_data = request.POST.copy()
        if 'nominee_first_name' in post_data and post_data['nominee_first_name']:
            full_name = post_data['nominee_first_name'].strip()
            name_parts = full_name.split(' ', 1)
            post_data['nominee_first_name'] = name_parts[0]
            if len(name_parts) > 1:
                post_data['nominee_last_name'] = name_parts[1]
            else:
                post_data['nominee_last_name'] = ''
        
        form = HallOfFameApplicationForm(post_data, request.FILES)
        if form.is_valid():
            application = form.save(commit=False)
            if request.user.is_authenticated:
                application.submitted_by = request.user
            application.save()
            messages.success(request, 'Your Hall of Fame application has been submitted successfully!')
            return redirect('app:hall_of_fame')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = HallOfFameApplicationForm()
    
    return render(request, 'app/hall_of_fame_apply.html', {'form': form})

@check_page_enabled('enable_check_turnitin_page')
def check_turnitin(request):
    """Check Turnitin Plagiarism page view"""
    if request.method == 'POST':
        form = PlagiarismCheckForm(request.POST, request.FILES)
        if form.is_valid():
            check = form.save(commit=False)
            if request.user.is_authenticated:
                check.submitted_by = request.user
            check.save()
            messages.success(request, 'Your document has been submitted for plagiarism check!')
            return redirect('app:landing')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = PlagiarismCheckForm()
    
    return render(request, 'app/check_turnitin.html', {'form': form})

@check_page_enabled('enable_work_plagiarism_page')
def work_plagiarism(request):
    """Work My Plagiarism page view"""
    if request.method == 'POST':
        form = PlagiarismWorkForm(request.POST, request.FILES)
        if form.is_valid():
            work = form.save(commit=False)
            if request.user.is_authenticated:
                work.submitted_by = request.user
            work.save()
            messages.success(request, 'Your request has been submitted successfully!')
            return redirect('app:landing')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = PlagiarismWorkForm()
    
    return render(request, 'app/work_plagiarism.html', {'form': form})

@check_page_enabled('enable_thesis_to_article_page')
def thesis_to_article(request):
    """Thesis To Article page view"""
    if request.method == 'POST':
        form = ThesisToArticleForm(request.POST, request.FILES)
        if form.is_valid():
            conversion = form.save(commit=False)
            if request.user.is_authenticated:
                conversion.submitted_by = request.user
            conversion.save()
            messages.success(request, 'Your thesis has been submitted for article conversion!')
            return redirect('app:landing')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = ThesisToArticleForm()
    
    return render(request, 'app/thesis_to_article.html', {'form': form})

@check_page_enabled('enable_thesis_to_book_page')
def thesis_to_book(request):
    """Thesis To Book page view"""
    if request.method == 'POST':
        form = ThesisToBookForm(request.POST, request.FILES)
        if form.is_valid():
            conversion = form.save(commit=False)
            if request.user.is_authenticated:
                conversion.submitted_by = request.user
            conversion.save()
            messages.success(request, 'Your thesis has been submitted for book conversion!')
            return redirect('app:landing')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = ThesisToBookForm()
    
    return render(request, 'app/thesis_to_book.html', {'form': form})

@check_page_enabled('enable_thesis_to_book_chapter_page')
def thesis_to_book_chapter(request):
    """Thesis To Book Chapter page view"""
    if request.method == 'POST':
        form = ThesisToBookChapterForm(request.POST, request.FILES)
        if form.is_valid():
            conversion = form.save(commit=False)
            if request.user.is_authenticated:
                conversion.submitted_by = request.user
            conversion.save()
            messages.success(request, 'Your thesis has been submitted for book chapter conversion!')
            return redirect('app:landing')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = ThesisToBookChapterForm()
    
    return render(request, 'app/thesis_to_book_chapter.html', {'form': form})

@check_page_enabled('enable_powerpoint_preparation_page')
def powerpoint_preparation(request):
    """Power Point Preparation page view"""
    if request.method == 'POST':
        form = PowerPointPreparationForm(request.POST, request.FILES)
        if form.is_valid():
            preparation = form.save(commit=False)
            if request.user.is_authenticated:
                preparation.submitted_by = request.user
            preparation.save()
            messages.success(request, 'Your thesis has been submitted for PowerPoint preparation!')
            return redirect('app:landing')
        else:
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
    else:
        form = PowerPointPreparationForm()
    
    return render(request, 'app/powerpoint_preparation.html', {'form': form})

@check_page_enabled('enable_about_sis_page')
def about_sis(request):
    """About S.I.S page view"""
    return render(request, 'app/about_sis.html')

@check_page_enabled('enable_mission_page')
def mission(request):
    """Mission page view"""
    return render(request, 'app/mission.html')

@check_page_enabled('enable_criteria_page')
def criteria(request):
    """Criteria page view"""
    return render(request, 'app/criteria.html')

@check_page_enabled('enable_tolerance_policy_page')
def tolerance_policy(request):
    """Tolerance Policy page view"""
    return render(request, 'app/tolerance_policy.html')

@check_page_enabled('enable_service_solution_page')
def service_solution(request):
    """Service & Solution page view"""
    return render(request, 'app/service_solution.html')

@check_page_enabled('enable_policy_terms_page')
def policy_terms(request):
    """Policy Terms and Conditions page view"""
    return render(request, 'app/policy_terms.html')

def news(request):
    """News page view with Latest News, Top Tags, and Recommended News"""
    # Get latest news articles (10 for the grid)
    latest_news = NewsArticle.objects.filter(is_published=True)[:10]
    
    # Get top tags (ordered by priority, then name, limit to 3 initially)
    top_tags = NewsTag.objects.filter(is_active=True).order_by('-order_priority', 'name')[:3]
    
    # Get recommended news articles (9 items)
    recommended_news = NewsArticle.objects.filter(is_published=True)[:9]
    
    return render(request, 'app/news.html', {
        'latest_news': latest_news,
        'top_tags': top_tags,
        'recommended_news': recommended_news,
    })

def load_more_tags(request):
    """AJAX endpoint to load more tags"""
    offset = int(request.GET.get('offset', 3))
    limit = int(request.GET.get('limit', 10))
    
    # Get tags with limit + 1 to check if there are more
    all_tags = NewsTag.objects.filter(is_active=True).order_by('-order_priority', 'name')
    total_count = all_tags.count()
    tags = all_tags[offset:offset + limit]
    
    tags_data = []
    for tag in tags:
        tags_data.append({
            'id': tag.id,
            'name': tag.name,
        })
    
    # Check if there are more tags available
    has_more = (offset + len(tags_data)) < total_count
    
    return JsonResponse({
        'success': True,
        'tags': tags_data,
        'has_more': has_more
    })

def browse_all_news(request):
    """Browse all news with search and filter functionality"""
    from django.db.models import Q
    from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
    
    # Get all published news articles
    articles = NewsArticle.objects.filter(is_published=True)
    
    # Search functionality
    search_query = request.GET.get('search', '').strip()
    if search_query:
        articles = articles.filter(
            Q(title__icontains=search_query) |
            Q(content__icontains=search_query) |
            Q(excerpt__icontains=search_query)
        )
    
    # Filter by tag
    tag_filter = request.GET.get('tag', '')
    selected_tag_obj = None
    if tag_filter:
        try:
            selected_tag_obj = NewsTag.objects.get(id=tag_filter, is_active=True)
            articles = articles.filter(tags__id=tag_filter)
        except NewsTag.DoesNotExist:
            tag_filter = ''
    
    # Filter by date range (optional - can be added later)
    # date_from = request.GET.get('date_from', '')
    # date_to = request.GET.get('date_to', '')
    
    # Order by published date (newest first)
    articles = articles.order_by('-published_date').distinct()
    
    # Get all active tags for filter dropdown
    all_tags = NewsTag.objects.filter(is_active=True).order_by('name')
    
    # Pagination - 15 articles per page
    paginator = Paginator(articles, 15)
    page = request.GET.get('page', 1)
    
    try:
        page_obj = paginator.get_page(page)
    except PageNotAnInteger:
        page_obj = paginator.get_page(1)
    except EmptyPage:
        page_obj = paginator.get_page(paginator.num_pages)
    
    return render(request, 'app/browse_all_news.html', {
        'articles': page_obj,
        'page_obj': page_obj,
        'all_tags': all_tags,
        'search_query': search_query,
        'selected_tag': tag_filter,
        'selected_tag_obj': selected_tag_obj,
    })

def hero_autocomplete(request):
    """Autocomplete API endpoint for hero search - searches only: Indexed Articles, Indexed Journals, Project Archive, and News"""
    from django.utils.html import strip_tags
    import re
    
    query = request.GET.get('q', '').strip()
    
    if len(query) < 2:
        return JsonResponse({'success': False, 'results': []})
    
    def clean_text(text):
        """Remove HTML tags and code snippets from text"""
        if not text:
            return ""
        # Remove HTML tags
        text = strip_tags(str(text))
        # Remove code-like patterns (HTML entities, style attributes, etc.)
        text = re.sub(r'&[a-z]+;', '', text)  # Remove HTML entities like &quot;
        text = re.sub(r'style="[^"]*"', '', text)  # Remove style attributes
        text = re.sub(r'<[^>]+>', '', text)  # Remove any remaining HTML tags
        text = re.sub(r'\{[^}]+\}', '', text)  # Remove CSS-like blocks
        text = re.sub(r'rgb\([^)]+\)', '', text)  # Remove rgb() color codes
        # Clean up extra whitespace
        text = ' '.join(text.split())
        return text[:100]  # Limit length
    
    results = []
    
    # 1. Search Indexed Articles
    articles = Article.objects.filter(
        status__in=['approved', 'pending']  # Only show approved/pending articles
    ).filter(
        Q(title__icontains=query) | 
        Q(abstract__icontains=query) | 
        Q(keywords__icontains=query) |
        Q(authors_names__icontains=query) |
        Q(journal_name__icontains=query) |
        Q(discipline__icontains=query)
    )[:8]  # Get more results per category
    for article in articles:
        results.append({
            'title': clean_text(article.title),
            'type': 'Indexed Articles',
            'url': f'/indexed_articles/view/{article.id}/',
            'requires_auth': False
        })
    
    # 2. Search Indexed Journals
    journals = Journal.objects.filter(
        Q(journal_name__icontains=query) | 
        Q(journal_scope__icontains=query) |
        Q(publisher_name__icontains=query) |
        Q(issn_print__icontains=query) |
        Q(issn_online__icontains=query) |
        Q(e_issn__icontains=query) |
        Q(subject_area__icontains=query)
    )[:8]
    for journal in journals:
        results.append({
            'title': clean_text(journal.journal_name),
            'type': 'Indexed Journals',
            'url': f'/indexed_journals/view/{journal.id}/',
            'requires_auth': False
        })
    
    # 3. Search Project | Research Archive
    projects = Project.objects.filter(
        Q(project_title__icontains=query) | 
        Q(description__icontains=query) |
        Q(category__icontains=query) |
        Q(institution__icontains=query)
    )[:8]
    for project in projects:
        results.append({
            'title': clean_text(project.project_title),
            'type': 'Project | Research Archive',
            'url': f'/project_archive/view/{project.id}/',
            'requires_auth': False
        })
    
    # 4. Search News Page
    # Search in title, excerpt, content, and tags
    query_lower = query.lower()
    
    # If query contains "news" or "article", also show some recent news articles
    if 'news' in query_lower or 'article' in query_lower:
        # Get matching articles first
        news_articles = NewsArticle.objects.filter(
            is_published=True
        ).filter(
            Q(title__icontains=query) | 
            Q(excerpt__icontains=query) |
            Q(content__icontains=query) |
            Q(tags__name__icontains=query)
        ).distinct()[:8]
        
        # If no matches but query contains "news", show recent articles
        if not news_articles.exists() and 'news' in query_lower:
            news_articles = NewsArticle.objects.filter(
                is_published=True
            ).order_by('-published_date')[:8]
    else:
        # Regular search
        news_articles = NewsArticle.objects.filter(
            is_published=True
        ).filter(
            Q(title__icontains=query) | 
            Q(excerpt__icontains=query) |
            Q(content__icontains=query) |
            Q(tags__name__icontains=query)
        ).distinct()[:8]
    
    for news in news_articles:
        results.append({
            'title': clean_text(news.title),
            'type': 'News',
            'url': f'/news/{news.slug}/',
            'requires_auth': False
        })
    
    # Limit total results and sort by relevance (prioritize title matches)
    # Sort: title matches first, then other matches
    results.sort(key=lambda x: (not x['title'].lower().startswith(query.lower()), x['title']))
    results = results[:20]  # Show up to 20 results total
    
    return JsonResponse({
        'success': True,
        'results': results
    })

@login_required
def create_news(request):
    """Create new news article page"""
    # Get or create NewsWriter for the logged-in user
    writer_name = request.user.get_full_name() or request.user.username
    writer, created = NewsWriter.objects.get_or_create(
        name=writer_name,
        defaults={
            'email': request.user.email or '',
            'is_active': True
        }
    )
    
    if request.method == 'POST':
        form = NewsArticleForm(request.POST, request.FILES)
        
        # Debug: Print form data
        print(f"Form data: {request.POST}")
        print(f"Content value: {request.POST.get('content', 'NOT SET')}")
        print(f"Title value: {request.POST.get('title', 'NOT SET')}")
        
        if form.is_valid():
            article = form.save(commit=False)
            article.created_by = request.user
            article.writer = writer  # Auto-set writer
            
            # Handle draft vs publish
            action = request.POST.get('action', 'publish')
            if action == 'draft':
                article.is_published = False
            else:
                article.is_published = form.cleaned_data.get('is_published', False)
            
            article.save()
            form.save_m2m()  # Save many-to-many relationships (tags)
            
            if action == 'draft':
                messages.success(request, 'News article saved as draft!')
            else:
                messages.success(request, 'News article created successfully!')
            return redirect('app:news')
        else:
            # Display all form errors
            print(f"Form errors: {form.errors}")
            print(f"Form non-field errors: {form.non_field_errors()}")
            for field, errors in form.errors.items():
                for error in errors:
                    messages.error(request, f'{field}: {error}')
            # Also show a general error message
            if not form.errors:
                messages.error(request, 'Please check the form and try again.')
    else:
        form = NewsArticleForm(initial={'writer': writer})
    
    # Get predefined tags or create them
    predefined_tags = [
        'Breaking News', 'Politics', 'Business', 'Technology', 'Sports',
        'Entertainment', 'Health', 'World News', 'Local News', 'Economy',
        'Crime', 'Education', 'Environment', 'Opinion'
    ]
    
    # Ensure all predefined tags exist
    for tag_name in predefined_tags:
        NewsTag.objects.get_or_create(
            name=tag_name,
            defaults={'is_active': True}
        )
    
    tags = NewsTag.objects.filter(is_active=True).order_by('name')
    
    return render(request, 'app/create_news.html', {
        'form': form,
        'tags': tags,
        'predefined_tags': predefined_tags
    })

def news_detail(request, slug):
    """News article detail page"""
    article = get_object_or_404(NewsArticle, slug=slug, is_published=True)
    
    # Increment view count
    article.view_count += 1
    article.save(update_fields=['view_count'])
    
    # Get sidebar articles (other published articles, excluding current)
    sidebar_articles = NewsArticle.objects.filter(
        is_published=True
    ).exclude(id=article.id).order_by('-published_date')[:5]
    
    # Get related articles (same tags, excluding current article)
    related_articles = NewsArticle.objects.filter(
        tags__in=article.tags.all(),
        is_published=True
    ).exclude(id=article.id).distinct()[:4]
    
    # Get approved comments (top-level only, replies are loaded via AJAX)
    comments = NewsComment.objects.filter(
        article=article,
        parent=None,
        is_approved=True
    ).order_by('-created_at')
    
    # Check if user has bookmarked this article
    is_bookmarked = False
    if request.user.is_authenticated:
        is_bookmarked = NewsBookmark.objects.filter(user=request.user, article=article).exists()
    
    return render(request, 'app/news_detail.html', {
        'article': article,
        'sidebar_articles': sidebar_articles,
        'related_articles': related_articles,
        'comments': comments,
        'is_bookmarked': is_bookmarked,
    })

@login_required
def add_news_comment(request, article_slug):
    """Add a comment to a news article"""
    if request.method == 'POST':
        article = get_object_or_404(NewsArticle, slug=article_slug, is_published=True)
        content = request.POST.get('content', '').strip()
        parent_id = request.POST.get('parent_id', None)
        
        if content:
            comment = NewsComment.objects.create(
                article=article,
                user=request.user,
                content=content,
                parent_id=parent_id if parent_id else None
            )
            messages.success(request, 'Your comment has been added!')
        else:
            messages.error(request, 'Comment cannot be empty.')
    
    return redirect('app:news_detail', slug=article_slug)

@login_required
def like_news_comment(request, comment_id):
    """Like or unlike a comment"""
    comment = get_object_or_404(NewsComment, id=comment_id, is_approved=True)
    
    if request.user in comment.likes.all():
        comment.likes.remove(request.user)
        action = 'unliked'
    else:
        comment.likes.add(request.user)
        comment.dislikes.remove(request.user)  # Remove from dislikes if present
        action = 'liked'
    
    return JsonResponse({
        'success': True,
        'action': action,
        'likes_count': comment.get_likes_count(),
        'dislikes_count': comment.get_dislikes_count()
    })

@login_required
def dislike_news_comment(request, comment_id):
    """Dislike or undislike a comment"""
    comment = get_object_or_404(NewsComment, id=comment_id, is_approved=True)
    
    if request.user in comment.dislikes.all():
        comment.dislikes.remove(request.user)
        action = 'undisliked'
    else:
        comment.dislikes.add(request.user)
        comment.likes.remove(request.user)  # Remove from likes if present
        action = 'disliked'
    
    return JsonResponse({
        'success': True,
        'action': action,
        'likes_count': comment.get_likes_count(),
        'dislikes_count': comment.get_dislikes_count()
    })

def get_comment_replies(request, comment_id):
    """Get replies for a comment (AJAX)"""
    comment = get_object_or_404(NewsComment, id=comment_id, is_approved=True)
    replies = comment.replies.filter(is_approved=True).order_by('created_at')
    
    replies_data = []
    for reply in replies:
        replies_data.append({
            'id': reply.id,
            'user': reply.user.username,
            'content': reply.content,
            'created_at': reply.created_at.strftime('%B %d, %Y at %I:%M %p'),
            'likes_count': reply.get_likes_count(),
            'dislikes_count': reply.get_dislikes_count(),
            'user_liked': request.user.is_authenticated and request.user in reply.likes.all(),
            'user_disliked': request.user.is_authenticated and request.user in reply.dislikes.all(),
        })
    
    return JsonResponse({'success': True, 'replies': replies_data})

@login_required
def toggle_bookmark(request, article_slug):
    """Toggle bookmark for a news article"""
    article = get_object_or_404(NewsArticle, slug=article_slug, is_published=True)
    
    bookmark, created = NewsBookmark.objects.get_or_create(
        user=request.user,
        article=article
    )
    
    if not created:
        # Bookmark already exists, remove it
        bookmark.delete()
        is_bookmarked = False
        action = 'removed'
    else:
        # Bookmark was created
        is_bookmarked = True
        action = 'added'
    
    return JsonResponse({
        'success': True,
        'action': action,
        'is_bookmarked': is_bookmarked
    })

@login_required
def delete_news_comment(request, comment_id):
    """Delete a comment (superadmin only)"""
    comment = get_object_or_404(NewsComment, id=comment_id)
    
    # Check if user is superuser (superadmin)
    if not request.user.is_superuser:
        return JsonResponse({'success': False, 'error': 'Permission denied. Only superadmins can delete comments.'}, status=403)
    
    article_slug = comment.article.slug
    comment.delete()
    messages.success(request, 'Comment deleted successfully.')
    
    if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
        return JsonResponse({'success': True})
    
    return redirect('app:news_detail', slug=article_slug)

@login_required
def news_list(request):
    """Published blogs/news listing page"""
    from app.models import Blog
    blogs = Blog.objects.filter(is_published=True).order_by('-published_date', '-order_priority')
    return render(request, 'app/news_list.html', {'blogs': blogs})

@login_required
def news_add(request):
    """Add new blog/news page"""
    from app.models import Blog
    from django.contrib import messages
    
    if request.method == 'POST':
        blog = Blog.objects.create(
            title=request.POST.get('blog_title', ''),
            content=request.POST.get('blog_content', ''),
            tag=request.POST.get('blog_tag', ''),
            category=request.POST.get('blog_category', 'general'),
            order_priority=int(request.POST.get('blog_order_priority', 5)),
            created_by=request.user,
        )
        if 'blog_image' in request.FILES:
            blog.image = request.FILES['blog_image']
        blog.save()
        messages.success(request, 'Blog created successfully!')
        return redirect('app:news_list')
    
    return render(request, 'app/news_add.html')

@login_required
def news_edit(request, blog_id):
    """Edit blog/news page"""
    from app.models import Blog
    from django.contrib import messages
    from django.shortcuts import get_object_or_404
    
    blog = get_object_or_404(Blog, id=blog_id)
    
    if request.method == 'POST':
        blog.title = request.POST.get('blog_title', '')
        blog.content = request.POST.get('blog_content', '')
        blog.tag = request.POST.get('blog_tag', '')
        blog.category = request.POST.get('blog_category', 'general')
        blog.order_priority = int(request.POST.get('blog_order_priority', 5))
        if 'blog_image' in request.FILES:
            blog.image = request.FILES['blog_image']
        blog.save()
        messages.success(request, 'Blog updated successfully!')
        return redirect('app:news_list')
    
    return render(request, 'app/news_edit.html', {'blog': blog})

@login_required
def news_delete(request, blog_id):
    """Delete blog/news"""
    from app.models import Blog
    from django.http import JsonResponse
    from django.shortcuts import get_object_or_404
    
    blog = get_object_or_404(Blog, id=blog_id)
    blog.delete()
    return JsonResponse({'success': True, 'message': 'Blog deleted successfully!'})

@login_required
def news_toggle_publish(request, blog_id):
    """Toggle blog publish status"""
    from app.models import Blog
    from django.http import JsonResponse
    from django.shortcuts import get_object_or_404
    
    blog = get_object_or_404(Blog, id=blog_id)
    blog.is_published = not blog.is_published
    blog.save()
    return JsonResponse({'success': True, 'message': 'Blog status updated successfully!'})

def dashboard(request):
    """Dashboard page view"""
    user = request.user
    profile = None
    if user.is_authenticated:
        try:
            profile = user.profile
        except UserProfile.DoesNotExist:
            profile = UserProfile.objects.create(user=user)
    
    # Get statistics
    articles_count = Article.objects.filter(status='approved').count()
    journals_count = Journal.objects.count()
    projects_count = Project.objects.count()
    
    # Get all pending requests from different models
    pending_plagiarism_checks = PlagiarismCheck.objects.all()
    pending_plagiarism_works = PlagiarismWork.objects.all()
    pending_thesis_articles = ThesisToArticle.objects.all()
    pending_thesis_books = ThesisToBook.objects.all()
    pending_thesis_chapters = ThesisToBookChapter.objects.all()
    pending_powerpoints = PowerPointPreparation.objects.all()
    
    # Combine all requests and get latest 10
    all_requests = []
    for req in pending_plagiarism_checks:
        all_requests.append({
            'id': req.id,
            'type': 'check_plagiarism',
            'email': req.email,
            'created_at': req.created_at,
            'document': req.document,
            'name': req.name or '',
        })
    for req in pending_plagiarism_works:
        all_requests.append({
            'id': req.id,
            'type': 'work_plagiarism',
            'email': req.email,
            'created_at': req.created_at,
            'document': req.document,
            'name': req.name or '',
        })
    for req in pending_thesis_articles:
        all_requests.append({
            'id': req.id,
            'type': 'thesis_to_article',
            'email': req.email,
            'created_at': req.created_at,
            'document': req.thesis_file,
            'name': req.name or '',
        })
    for req in pending_thesis_books:
        all_requests.append({
            'id': req.id,
            'type': 'thesis_to_book',
            'email': req.email,
            'created_at': req.created_at,
            'document': req.thesis_file,
            'name': req.name or '',
        })
    for req in pending_thesis_chapters:
        all_requests.append({
            'id': req.id,
            'type': 'thesis_to_book_chapter',
            'email': req.email,
            'created_at': req.created_at,
            'document': req.thesis_file,
            'name': req.name or '',
        })
    for req in pending_powerpoints:
        all_requests.append({
            'id': req.id,
            'type': 'powerpoint_preparation',
            'email': req.email,
            'created_at': req.created_at,
            'document': req.thesis_file,
            'name': req.name or '',
        })
    
    # Sort by created_at and get latest 10
    all_requests.sort(key=lambda x: x['created_at'] if x['created_at'] else None, reverse=True)
    latest_requests = all_requests[:10]
    pending_requests_count = len(all_requests)
    
    # Get account statistics
    total_users = User.objects.count()
    directory_researchers_count = DirectoryApplication.objects.filter(terms_accepted=True).count()
    eminent_personalities_count = HallOfFameApplication.objects.count()
    registered_members_count = MembershipRequest.objects.count()
    
    # Get recent accounts (users with profiles)
    recent_users = User.objects.select_related('profile').order_by('-date_joined')[:5]
    
    # Get recent blog posts (using articles as blog posts for now)
    recent_blogs = Article.objects.filter(status='approved').order_by('-created_at')[:5]
    
    return render(request, 'app/dashboard.html', {
        'user': user,
        'profile': profile,
        'articles_count': articles_count,
        'journals_count': journals_count,
        'projects_count': projects_count,
        'pending_requests_count': pending_requests_count,
        'latest_requests': latest_requests,
        'total_users': total_users,
        'directory_researchers_count': directory_researchers_count,
        'eminent_personalities_count': eminent_personalities_count,
        'registered_members_count': registered_members_count,
        'recent_users': recent_users,
        'recent_blogs': recent_blogs,
    })

@login_required
def user_dashboard(request):
    """User Dashboard page view - for ordinary users"""
    user = request.user
    profile = None
    try:
        profile = user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=user)
    
    # Get user's articles count
    user_articles_count = Article.objects.filter(submitted_by=user).count()
    
    # Get user's journals count
    user_journals_count = Journal.objects.filter(submitted_by=user).count()
    
    # Get user's requests count (all types)
    user_plagiarism_checks = PlagiarismCheck.objects.filter(email=user.email)
    user_plagiarism_works = PlagiarismWork.objects.filter(email=user.email)
    user_thesis_articles = ThesisToArticle.objects.filter(email=user.email)
    user_thesis_books = ThesisToBook.objects.filter(email=user.email)
    user_thesis_chapters = ThesisToBookChapter.objects.filter(email=user.email)
    user_powerpoints = PowerPointPreparation.objects.filter(email=user.email)
    
    user_requests_count = (
        user_plagiarism_checks.count() +
        user_plagiarism_works.count() +
        user_thesis_articles.count() +
        user_thesis_books.count() +
        user_thesis_chapters.count() +
        user_powerpoints.count()
    )
    
    # Get user's articles
    user_articles = Article.objects.filter(submitted_by=user).order_by('-created_at')
    
    # Get user's journals
    user_journals = Journal.objects.filter(submitted_by=user).order_by('-created_at')
    
    # Get user's requests (combined)
    import hashlib
    user_requests = []
    for req in user_plagiarism_checks:
        # Generate unique ID from request
        req_id = hashlib.md5(f"check_plagiarism_{req.id}_{req.created_at}".encode()).hexdigest()[:12]
        description = req.document_title or f"Document {req.id}"
        user_requests.append({
            'id': req.id,
            'type': 'check_plagiarism',
            'type_display': 'Check Plagiarism',
            'description': description,
            'unique_id': req_id.upper(),
            'created_at': req.created_at,
            'status': 'declined' if not req.terms_accepted else 'pending',
            'has_file': bool(req.document),
        })
    for req in user_plagiarism_works:
        req_id = hashlib.md5(f"work_plagiarism_{req.id}_{req.created_at}".encode()).hexdigest()[:12]
        description = req.submission_title or f"Work {req.id}"
        user_requests.append({
            'id': req.id,
            'type': 'work_plagiarism',
            'type_display': 'Work Plagiarism',
            'description': description,
            'unique_id': req_id.upper(),
            'created_at': req.created_at,
            'status': 'declined' if not req.terms_accepted else 'pending',
            'has_file': bool(req.document),
        })
    for req in user_thesis_articles:
        req_id = hashlib.md5(f"thesis_to_article_{req.id}_{req.created_at}".encode()).hexdigest()[:12]
        description = req.submission_title or f"Thesis {req.id}"
        user_requests.append({
            'id': req.id,
            'type': 'thesis_to_article',
            'type_display': 'Thesis To Article',
            'description': description,
            'unique_id': req_id.upper(),
            'created_at': req.created_at,
            'status': 'declined' if not req.terms_accepted else 'pending',
            'has_file': bool(req.thesis_file),
        })
    for req in user_thesis_books:
        req_id = hashlib.md5(f"thesis_to_book_{req.id}_{req.created_at}".encode()).hexdigest()[:12]
        description = req.submission_title or req.book_title or f"Book {req.id}"
        user_requests.append({
            'id': req.id,
            'type': 'thesis_to_book',
            'type_display': 'Thesis To Book',
            'description': description,
            'unique_id': req_id.upper(),
            'created_at': req.created_at,
            'status': 'declined' if not req.terms_accepted else 'pending',
            'has_file': bool(req.thesis_file),
        })
    for req in user_thesis_chapters:
        req_id = hashlib.md5(f"thesis_to_book_chapter_{req.id}_{req.created_at}".encode()).hexdigest()[:12]
        description = req.submission_title or req.chapter_title or f"Chapter {req.id}"
        user_requests.append({
            'id': req.id,
            'type': 'thesis_to_book_chapter',
            'type_display': 'Thesis To Book Chapter',
            'description': description,
            'unique_id': req_id.upper(),
            'created_at': req.created_at,
            'status': 'declined' if not req.terms_accepted else 'pending',
            'has_file': bool(req.thesis_file),
        })
    for req in user_powerpoints:
        req_id = hashlib.md5(f"powerpoint_preparation_{req.id}_{req.created_at}".encode()).hexdigest()[:12]
        description = req.submission_title or f"Presentation {req.id}"
        user_requests.append({
            'id': req.id,
            'type': 'powerpoint_preparation',
            'type_display': 'Power Point Preparation',
            'description': description,
            'unique_id': req_id.upper(),
            'created_at': req.created_at,
            'status': 'declined' if not req.terms_accepted else 'pending',
            'has_file': bool(req.thesis_file),
        })
    
    # Sort requests by created_at
    user_requests.sort(key=lambda x: x['created_at'] if x['created_at'] else None, reverse=True)
    
    # Get user's bookmarks
    user_bookmarks = NewsBookmark.objects.filter(user=user).select_related('article').prefetch_related('article__tags').order_by('-created_at')
    
    # Get user's news articles
    user_news_articles = NewsArticle.objects.filter(created_by=user).order_by('-created_at')
    user_news_articles_count = user_news_articles.count()
    
    return render(request, 'app/user_dashboard.html', {
        'user': user,
        'profile': profile,
        'user_articles_count': user_articles_count,
        'user_journals_count': user_journals_count,
        'user_requests_count': user_requests_count,
        'user_articles': user_articles,
        'user_journals': user_journals,
        'user_requests': user_requests,
        'user_bookmarks': user_bookmarks,
        'user_news_articles': user_news_articles,
        'user_news_articles_count': user_news_articles_count,
    })

@login_required
def settings(request):
    """Settings page view"""
    user = request.user
    try:
        profile = user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=user)
    
    if request.method == 'POST':
        section = request.POST.get('section', 'profile')
        
        if section == 'profile':
            # Update profile information
            user.first_name = request.POST.get('first_name', '')
            user.last_name = request.POST.get('last_name', '')
            user.email = request.POST.get('email', '')
            user.save()
            
            profile.phone = request.POST.get('phone', '')
            profile.country = request.POST.get('country', '')
            if 'profile_photo' in request.FILES:
                profile.profile_photo = request.FILES['profile_photo']
            profile.save()
            
            messages.success(request, 'Profile updated successfully!')
            
        elif section == 'account':
            # Update account settings
            username = request.POST.get('username', '')
            if username and username != user.username:
                if User.objects.filter(username=username).exclude(id=user.id).exists():
                    messages.error(request, 'Username already exists.')
                else:
                    user.username = username
                    user.save()
                    messages.success(request, 'Account settings updated successfully!')
            else:
                messages.success(request, 'Account settings updated successfully!')
                
        elif section == 'security':
            # Update password
            current_password = request.POST.get('current_password', '')
            new_password = request.POST.get('new_password', '')
            confirm_password = request.POST.get('confirm_password', '')
            
            if new_password:
                if not user.check_password(current_password):
                    messages.error(request, 'Current password is incorrect.')
                elif new_password != confirm_password:
                    messages.error(request, 'New passwords do not match.')
                elif len(new_password) < 8:
                    messages.error(request, 'Password must be at least 8 characters long.')
                else:
                    user.set_password(new_password)
                    user.save()
                    from django.contrib.auth import update_session_auth_hash
                    update_session_auth_hash(request, user)
                    messages.success(request, 'Password updated successfully!')
        
        elif section == 'general':
            # Update general settings - Theme
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            font_size = request.POST.get('font_size')
            if font_size:
                try:
                    site_settings.default_app_font_size = int(font_size)
                except ValueError:
                    pass
            
            site_settings.save()
            messages.success(request, 'General settings updated successfully!')
        
        elif section == 'carousel_1':
            # Update carousel image 1
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            if 'carousel_image_1' in request.FILES:
                site_settings.carousel_image_1 = request.FILES['carousel_image_1']
            
            carousel_image_1_url = request.POST.get('carousel_image_1_url', '').strip()
            if carousel_image_1_url:
                site_settings.carousel_image_1_url = carousel_image_1_url
            
            site_settings.carousel_image_1_title = request.POST.get('carousel_image_1_title', '')
            site_settings.carousel_image_1_subtitle = request.POST.get('carousel_image_1_subtitle', '')
            
            site_settings.save()
            messages.success(request, 'Carousel image 1 updated successfully!')
        
        elif section == 'carousel_2':
            # Update carousel image 2
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            if 'carousel_image_2' in request.FILES:
                site_settings.carousel_image_2 = request.FILES['carousel_image_2']
            
            carousel_image_2_url = request.POST.get('carousel_image_2_url', '').strip()
            if carousel_image_2_url:
                site_settings.carousel_image_2_url = carousel_image_2_url
            
            site_settings.carousel_image_2_title = request.POST.get('carousel_image_2_title', '')
            site_settings.carousel_image_2_subtitle = request.POST.get('carousel_image_2_subtitle', '')
            
            site_settings.save()
            messages.success(request, 'Carousel image 2 updated successfully!')
        
        elif section == 'carousel_3':
            # Update carousel image 3
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            if 'carousel_image_3' in request.FILES:
                site_settings.carousel_image_3 = request.FILES['carousel_image_3']
            
            carousel_image_3_url = request.POST.get('carousel_image_3_url', '').strip()
            if carousel_image_3_url:
                site_settings.carousel_image_3_url = carousel_image_3_url
            
            site_settings.carousel_image_3_title = request.POST.get('carousel_image_3_title', '')
            site_settings.carousel_image_3_subtitle = request.POST.get('carousel_image_3_subtitle', '')
            
            site_settings.save()
            messages.success(request, 'Carousel image 3 updated successfully!')
            
        elif section == 'address':
            # Update address & contact settings
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            site_settings.phone = request.POST.get('phone', '')
            site_settings.office_location_address = request.POST.get('office_address', '')
            site_settings.email_address = request.POST.get('email_address', '')
            site_settings.whatsapp_url = request.POST.get('whatsapp', '')
            site_settings.youtube_url = request.POST.get('youtube', '')
            site_settings.twitter_url = request.POST.get('twitter', '')
            site_settings.facebook_url = request.POST.get('facebook', '')
            
            site_settings.save()
            messages.success(request, 'Address & Contact settings updated successfully!')
            
        elif section == 'payment_gate':
            # Update payment gate settings
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            site_settings.api_secret_key = request.POST.get('api_secret_key', '')
            site_settings.api_public_key = request.POST.get('api_public_key', '')
            
            site_settings.save()
            messages.success(request, 'Payment gate settings updated successfully!')
            
        elif section == 'map_address':
            # Update map address
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            site_settings.map_index = request.POST.get('map_index', '')
            
            site_settings.save()
            messages.success(request, 'Map address updated successfully!')
            
        elif section == 'sponsorship_pricing':
            # Update sponsorship pricing
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            # Update all sponsorship prices
            for field in ['premier', 'sustaining', 'basic', 'power', 'double', 'single', 'compact', 'inspired', 'charity_homes']:
                price = request.POST.get(f'{field}_price', '')
                if price:
                    try:
                        setattr(site_settings, f'{field}_price', float(price))
                    except ValueError:
                        pass
            
            site_settings.save()
            messages.success(request, 'Sponsorship pricing updated successfully!')
            
        elif section == 'registrations_pricing':
            # Update registrations pricing
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            # Update all registration prices
            price_fields = {
                'article_indexing_price': 'article_indexing_price',
                'journal_indexing_price': 'journal_indexing_price',
                'project_archive_hosting_price': 'project_archive_hosting_price',
                'directory_researcher_price': 'directory_researcher_price',
                'hall_of_fame_price': 'hall_of_fame_price',
                'check_plagiarism_price': 'check_plagiarism_price',
                'work_plagiarism_price': 'work_plagiarism_price',
                'thesis_to_article_price': 'thesis_to_article_price',
                'thesis_to_book_price': 'thesis_to_book_price',
                'thesis_to_book_chapter_price': 'thesis_to_book_chapter_price',
                'powerpoint_preparation_price': 'powerpoint_preparation_price',
            }
            
            for form_field, model_field in price_fields.items():
                price = request.POST.get(form_field, '')
                if price:
                    try:
                        setattr(site_settings, model_field, float(price))
                    except ValueError:
                        pass
            
            site_settings.save()
            messages.success(request, 'Registrations pricing updated successfully!')
            
        elif section == 'navigators_documents':
            # Update navigators - documents
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            if 'documentations_nav_image' in request.FILES:
                site_settings.documentations_nav_image = request.FILES['documentations_nav_image']
            site_settings.documentations_nav_title = request.POST.get('documentations_nav_title', '')
            site_settings.documentations_nav_subtitle = request.POST.get('documentations_nav_subtitle', '')
            
            site_settings.save()
            messages.success(request, 'Documents navigator settings updated successfully!')
            
        elif section == 'navigators_communities':
            # Update navigators - communities
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            if 'communities_nav_image' in request.FILES:
                site_settings.communities_nav_image = request.FILES['communities_nav_image']
            site_settings.communities_nav_title = request.POST.get('communities_nav_title', '')
            site_settings.communities_nav_subtitle = request.POST.get('communities_nav_subtitle', '')
            
            site_settings.save()
            messages.success(request, 'Communities navigator settings updated successfully!')
            
        elif section == 'navigators_requests':
            # Update navigators - requests
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            if 'requests_nav_image' in request.FILES:
                site_settings.requests_nav_image = request.FILES['requests_nav_image']
            site_settings.requests_nav_title = request.POST.get('requests_nav_title', '')
            site_settings.requests_nav_subtitle = request.POST.get('requests_nav_subtitle', '')
            
            site_settings.save()
            messages.success(request, 'Requests navigator settings updated successfully!')
            
        elif section == 'navigators_about':
            # Update navigators - about
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            if 'about_nav_image' in request.FILES:
                site_settings.about_nav_image = request.FILES['about_nav_image']
            site_settings.about_nav_title = request.POST.get('about_nav_title', '')
            site_settings.about_nav_subtitle = request.POST.get('about_nav_subtitle', '')
            
            site_settings.save()
            messages.success(request, 'About navigator settings updated successfully!')
            
        elif section == 'landing_page':
            # Update landing page settings
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            site_settings.banner_title = request.POST.get('banner_title', '')
            site_settings.banner_subtitle = request.POST.get('banner_subtitle', '')
            site_settings.footer_video = request.POST.get('footer_video', '')
            site_settings.enable_service_features = request.POST.get('enable_service_features') == 'enabled'
            
            site_settings.save()
            messages.success(request, 'Landing page settings updated successfully!')
        
        elif section == 'toggle_page':
            # Toggle page enable/disable
            from app.models import SiteSettings
            site_settings = SiteSettings.get_settings()
            
            page_name = request.POST.get('page_name', '')
            is_enabled = request.POST.get('is_enabled', 'false') == 'true'
            
            page_mapping = {
                'landing': 'enable_landing_page',
                'indexed_articles': 'enable_indexed_articles_page',
                'indexed_journals': 'enable_indexed_journals_page',
                'project_archive': 'enable_project_archive_page',
                'directory_researchers': 'enable_directory_researchers_page',
                'hall_of_fame': 'enable_hall_of_fame_page',
                'council_members': 'enable_council_members_page',
                'team_members': 'enable_team_members_page',
                'donate': 'enable_donate_page',
                'about_sis': 'enable_about_sis_page',
                'mission': 'enable_mission_page',
                'criteria': 'enable_criteria_page',
                'tolerance_policy': 'enable_tolerance_policy_page',
                'service_solution': 'enable_service_solution_page',
                'policy_terms': 'enable_policy_terms_page',
                'check_turnitin': 'enable_check_turnitin_page',
                'work_plagiarism': 'enable_work_plagiarism_page',
                'thesis_to_article': 'enable_thesis_to_article_page',
                'thesis_to_book': 'enable_thesis_to_book_page',
                'thesis_to_book_chapter': 'enable_thesis_to_book_chapter_page',
                'powerpoint_preparation': 'enable_powerpoint_preparation_page',
            }
            
            if page_name in page_mapping:
                setattr(site_settings, page_mapping[page_name], is_enabled)
                site_settings.save()
                return JsonResponse({'success': True, 'message': f'Page {page_name} {"enabled" if is_enabled else "disabled"} successfully!'})
            else:
                return JsonResponse({'success': False, 'message': 'Invalid page name'})
        
        return redirect('app:settings')
    
    # Get all users for permissions management
    all_users = User.objects.all().order_by('username')
    # For now, we'll use a simple approach - in a real app, you'd have a Permission model
    users_with_permissions = all_users[:10]  # Limit to first 10 for display
    
    # Get site settings
    from app.models import SiteSettings
    site_settings = SiteSettings.get_settings()
    
    return render(request, 'app/settings.html', {
        'user': user,
        'profile': profile,
        'all_users': all_users,
        'users_with_permissions': users_with_permissions,
        'site_settings': site_settings,
    })

@login_required
def account_detail(request):
    """Account Detail page view"""
    user = request.user
    try:
        profile = user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=user)
    
    return render(request, 'app/account_detail.html', {
        'user': user,
        'profile': profile,
        'active_page': 'account_detail',
    })

@login_required
def list_accounts(request):
    """List all accounts page"""
    # Get search and filter parameters
    search_query = request.GET.get('search', '').strip()
    sort_by = request.GET.get('sort', 'recent')
    role_filter = request.GET.get('role', '')
    
    # Get all users with profiles
    users = User.objects.filter(profile__isnull=False).select_related('profile')
    
    # Get total count before filtering (for display)
    total_all_users = users.count()
    
    # Apply role filter
    if role_filter == 'admin':
        users = users.filter(is_superuser=True)
    elif role_filter == 'staff':
        users = users.filter(is_staff=True, is_superuser=False)
    elif role_filter == 'user':
        users = users.filter(is_staff=False, is_superuser=False)
    
    # Apply search filter
    if search_query:
        users = users.filter(
            Q(username__icontains=search_query) |
            Q(email__icontains=search_query) |
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query)
        )
    
    # Apply sorting
    if sort_by == 'recent':
        users = users.order_by('-date_joined')
    elif sort_by == 'oldest':
        users = users.order_by('date_joined')
    elif sort_by == 'name':
        users = users.order_by('first_name', 'last_name', 'username')
    else:
        users = users.order_by('-date_joined')
    
    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(users, 20)  # Show 20 users per page
    page_number = request.GET.get('page', 1)
    page_obj = paginator.get_page(page_number)
    
    # Calculate record range
    start_record = (page_obj.number - 1) * paginator.per_page + 1
    end_record = min(start_record + paginator.per_page - 1, paginator.count)
    
    # Ensure superusers have correct flags (fix any data inconsistencies)
    # Superusers should have both is_superuser=True and is_staff=True
    # Also check for users who might be superusers but have wrong flags
    for user in page_obj:
        if user.is_superuser:
            if not user.is_staff:
                user.is_staff = True
                user.save(update_fields=['is_staff'])
        # If user has is_staff=True but is_superuser=False, and they're the first superuser
        # or have a specific username pattern, we might want to check
        # For now, we'll let the frontend handle the display fix
    
    return render(request, 'app/list_accounts.html', {
        'page_obj': page_obj,
        'users': page_obj,
        'current_search': search_query,
        'current_sort': sort_by,
        'current_role': role_filter,
        'start_record': start_record,
        'end_record': end_record,
        'total_records': paginator.count,
        'total_all_users': total_all_users,
        'has_filters': bool(search_query or role_filter),
    })

@login_required
def view_account(request, user_id):
    """View Account page - displays account details for a specific user"""
    try:
        view_user = User.objects.get(id=user_id)
    except User.DoesNotExist:
        messages.error(request, 'User not found.')
        return redirect('app:dashboard')
    
    try:
        profile = view_user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=view_user)
    
    # Calculate document counts
    doc_counts = (
        Article.objects.filter(submitted_by=view_user).count() +
        Journal.objects.filter(submitted_by=view_user).count() +
        Project.objects.filter(submitted_by=view_user).count() +
        PlagiarismCheck.objects.filter(submitted_by=view_user).count() +
        PlagiarismWork.objects.filter(submitted_by=view_user).count() +
        ThesisToArticle.objects.filter(submitted_by=view_user).count() +
        ThesisToBook.objects.filter(submitted_by=view_user).count() +
        ThesisToBookChapter.objects.filter(submitted_by=view_user).count() +
        PowerPointPreparation.objects.filter(submitted_by=view_user).count()
    )
    
    # Get position from DirectoryApplication if exists
    position = None
    organization_url = None
    try:
        directory_app = DirectoryApplication.objects.filter(submitted_by=view_user, terms_accepted=True).first()
        if directory_app:
            position = directory_app.position
            # DirectoryApplication doesn't have organization_url, so we'll use institution or None
            organization_url = getattr(directory_app, 'institution', None)
    except:
        pass
    
    # Determine role (simplified - you can enhance this based on your needs)
    role = 'user'
    if view_user.is_staff:
        role = 'staff'
    if view_user.is_superuser:
        role = 'admin'
    
    # Status
    status = 'active' if view_user.is_active else 'inactive'
    
    return render(request, 'app/view_account.html', {
        'view_user': view_user,
        'profile': profile,
        'doc_counts': doc_counts,
        'position': position or 'N/A',
        'organization_url': organization_url,
        'role': role,
        'status': status,
    })

@login_required
def export_account(request, user_id, format_type):
    """Export account details in various formats"""
    try:
        view_user = User.objects.get(id=user_id)
    except User.DoesNotExist:
        messages.error(request, 'User not found.')
        return redirect('app:dashboard')
    
    try:
        profile = view_user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=view_user)
    
    # Calculate document counts
    doc_counts = (
        Article.objects.filter(submitted_by=view_user).count() +
        Journal.objects.filter(submitted_by=view_user).count() +
        Project.objects.filter(submitted_by=view_user).count() +
        PlagiarismCheck.objects.filter(submitted_by=view_user).count() +
        PlagiarismWork.objects.filter(submitted_by=view_user).count() +
        ThesisToArticle.objects.filter(submitted_by=view_user).count() +
        ThesisToBook.objects.filter(submitted_by=view_user).count() +
        ThesisToBookChapter.objects.filter(submitted_by=view_user).count() +
        PowerPointPreparation.objects.filter(submitted_by=view_user).count()
    )
    
    # Get position from DirectoryApplication if exists
    position = None
    organization_url = None
    try:
        directory_app = DirectoryApplication.objects.filter(submitted_by=view_user, terms_accepted=True).first()
        if directory_app:
            position = directory_app.position
            organization_url = getattr(directory_app, 'institution', None)
    except:
        pass
    
    # Determine role
    role = 'user'
    if view_user.is_staff:
        role = 'staff'
    if view_user.is_superuser:
        role = 'admin'
    
    status = 'active' if view_user.is_active else 'inactive'
    
    # Prepare account data
    account_data = {
        'Id': view_user.id,
        'Name': view_user.get_full_name() or view_user.username,
        'Email': view_user.email or 'N/A',
        'Phone': profile.phone or 'N/A',
        'Country': profile.country or 'N/A',
        'Status': status,
        'Role': role,
        'Doc Counts': doc_counts,
        'Itp': '0.5',
        'Ctrl': '97',
        'Profile Photo': profile.profile_photo.url if profile.profile_photo else 'null',
        'Position': position or 'N/A',
        'Organization Url': organization_url or 'null',
    }
    
    filename = f"account_{view_user.username}_{view_user.id}"
    
    if format_type == 'print':
        # Print-friendly HTML
        html_content = render_to_string('app/export_account_print.html', {
            'view_user': view_user,
            'account_data': account_data,
        })
        response = HttpResponse(html_content, content_type='text/html')
        return response
    
    elif format_type == 'pdf':
        # PDF export using reportlab or weasyprint
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.lib import colors
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
            from reportlab.lib.units import inch
            
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            elements = []
            
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                textColor=colors.HexColor('#111827'),
                spaceAfter=30,
            )
            
            elements.append(Paragraph("Account Details", title_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Create table data
            table_data = [['Field', 'Value']]
            for key, value in account_data.items():
                table_data.append([key, str(value)])
            
            table = Table(table_data, colWidths=[2*inch, 4*inch])
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#f9fafb')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.HexColor('#374151')),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 10),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.white),
                ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#111827')),
                ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb')),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f9fafb')]),
            ]))
            
            elements.append(table)
            doc.build(elements)
            
            response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
            return response
        except ImportError:
            # Fallback to simple text if reportlab not available
            response = HttpResponse(content_type='text/plain')
            response['Content-Disposition'] = f'attachment; filename="{filename}.txt"'
            for key, value in account_data.items():
                response.write(f"{key}: {value}\n")
            return response
    
    elif format_type == 'word':
        # Word export using python-docx
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            doc.add_heading('Account Details', 0)
            
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            
            # Header row
            header_cells = table.rows[0].cells
            header_cells[0].text = 'Field'
            header_cells[1].text = 'Value'
            for cell in header_cells:
                cell.paragraphs[0].runs[0].font.bold = True
            
            # Data rows
            for key, value in account_data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = key
                row_cells[1].text = str(value)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{filename}.docx"'
            return response
        except ImportError:
            # Fallback to simple text
            response = HttpResponse(content_type='text/plain')
            response['Content-Disposition'] = f'attachment; filename="{filename}.txt"'
            for key, value in account_data.items():
                response.write(f"{key}: {value}\n")
            return response
    
    elif format_type == 'csv':
        # CSV export
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['Field', 'Value'])
        for key, value in account_data.items():
            writer.writerow([key, value])
        
        return response
    
    elif format_type == 'excel':
        # Excel export using openpyxl
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment
            
            wb = Workbook()
            ws = wb.active
            ws.title = "Account Details"
            
            # Header row
            header_fill = PatternFill(start_color="F9FAFB", end_color="F9FAFB", fill_type="solid")
            header_font = Font(bold=True, color="374151")
            
            ws['A1'] = 'Field'
            ws['B1'] = 'Value'
            ws['A1'].fill = header_fill
            ws['B1'].fill = header_fill
            ws['A1'].font = header_font
            ws['B1'].font = header_font
            
            # Data rows
            row = 2
            for key, value in account_data.items():
                ws[f'A{row}'] = key
                ws[f'B{row}'] = str(value)
                if row % 2 == 0:
                    ws[f'A{row}'].fill = PatternFill(start_color="F9FAFB", end_color="F9FAFB", fill_type="solid")
                    ws[f'B{row}'].fill = PatternFill(start_color="F9FAFB", end_color="F9FAFB", fill_type="solid")
                row += 1
            
            # Adjust column widths
            ws.column_dimensions['A'].width = 20
            ws.column_dimensions['B'].width = 40
            
            buffer = io.BytesIO()
            wb.save(buffer)
            buffer.seek(0)
            
            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
            response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'
            return response
        except ImportError:
            # Fallback to CSV if openpyxl not available
            response = HttpResponse(content_type='text/csv')
            response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
            writer = csv.writer(response)
            writer.writerow(['Field', 'Value'])
            for key, value in account_data.items():
                writer.writerow([key, value])
            return response
    
    else:
        messages.error(request, 'Invalid export format.')
        return redirect('app:view_account', user_id=user_id)

@login_required
def update_user_role(request, user_id):
    """Update user role (Admin, Staff, User)"""
    # Superusers can change any role, staff can change to user/staff (not admin)
    if not (request.user.is_superuser or request.user.is_staff):
        return JsonResponse({'success': False, 'error': 'Permission denied. Only staff/superusers can change roles.'}, status=403)
    
    try:
        user = User.objects.get(id=user_id)
    except User.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'User not found'}, status=404)
    
    if request.method == 'POST':
        import json
        data = json.loads(request.body)
        role_type = data.get('role_type')
        
        # Set role based on selection
        # Only superusers can set admin role
        if role_type == 'admin':
            if not request.user.is_superuser:
                return JsonResponse({'success': False, 'error': 'Only superusers can assign admin role.'}, status=403)
            user.is_superuser = True
            user.is_staff = True  # Admin is also staff
            user.is_active = True
        elif role_type == 'staff':
            user.is_superuser = False
            user.is_staff = True
            user.is_active = True
        elif role_type == 'user':
            user.is_superuser = False
            user.is_staff = False
            # Keep is_active as is (don't deactivate when changing to user)
        
        # Save the user with explicit field updates to ensure persistence
        # Use update_fields to only update the specific fields we changed
        user.save(update_fields=['is_superuser', 'is_staff', 'is_active'])
        
        # Force database commit and verify the save worked
        from django.db import connection
        connection.ensure_connection()
        
        # Refresh from database to ensure we have the latest state
        user.refresh_from_db()
        
        # Verify the save worked
        print(f'User {user.id} role updated - is_superuser: {user.is_superuser}, is_staff: {user.is_staff}')
        
        # Send email notification to the user
        try:
            role_display = {
                'admin': 'Administrator',
                'staff': 'Staff',
                'user': 'User'
            }.get(role_type, role_type.capitalize())
            
            subject = f'Your Account Role Has Been Updated - ScholarIndex'
            
            # Use HTML email template if available, otherwise plain text
            try:
                html_message = render_to_string('app/emails/role_change_notification.html', {
                    'user': user,
                    'new_role_display': role_display,
                    'changed_by': request.user,
                })
                send_mail(
                    subject,
                    '',  # Plain text version (empty, using HTML)
                    settings.DEFAULT_FROM_EMAIL,
                    [user.email],
                    fail_silently=False,
                    html_message=html_message,
                )
            except Exception as template_error:
                # Fallback to plain text if template fails
                import logging
                logger = logging.getLogger(__name__)
                logger.warning(f'Could not load HTML email template: {str(template_error)}')
                
                message = f'''
Hello {user.get_full_name() or user.username},

Your account role on ScholarIndex has been updated.

New Role: {role_display}

This change was made by: {request.user.get_full_name() or request.user.username}

If you did not request this change or have any concerns, please contact our support team immediately.

Best regards,
ScholarIndex Team
'''
                send_mail(
                    subject,
                    message,
                    settings.DEFAULT_FROM_EMAIL,
                    [user.email],
                    fail_silently=False,
                )
        except Exception as e:
            # Log the error but don't fail the role update
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f'Failed to send role change email to {user.email}: {str(e)}', exc_info=True)
            # Print to console for debugging
            print(f'Email send error: {str(e)}')
            # Still return success since role was updated, but include email error in response
            return JsonResponse({
                'success': True,
                'message': f'Role updated to {role_type} successfully',
                'is_superuser': user.is_superuser,
                'is_staff': user.is_staff,
                'email_sent': False,
                'email_error': str(e)
            })
        
        return JsonResponse({
            'success': True,
            'message': f'Role updated to {role_type} successfully',
            'is_superuser': user.is_superuser,
            'is_staff': user.is_staff,
            'email_sent': True,
        })
    
    return JsonResponse({'success': False, 'error': 'Invalid request'}, status=400)

@login_required
def get_user_data(request, user_id):
    """Get user data for edit modal"""
    # Allow any logged-in user to view account data (they can see it in the list anyway)
    
    try:
        user = User.objects.get(id=user_id)
    except User.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'User not found'}, status=404)
    
    try:
        profile = user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=user)
    
    # Determine role
    role = 'user'
    if user.is_superuser:
        role = 'admin'
    elif user.is_staff:
        role = 'staff'
    
    return JsonResponse({
        'success': True,
        'user': {
            'username': user.username or '',
            'email': user.email or '',
            'first_name': user.first_name or '',
            'last_name': user.last_name or '',
            'is_active': user.is_active,
            'is_superuser': user.is_superuser,
            'is_staff': user.is_staff,
            'role': role,
        },
        'profile': {
            'phone': profile.phone or '',
            'country': profile.country or '',
        }
    })

@login_required
def edit_user_account(request, user_id):
    """Edit user account page"""
    # Allow any logged-in user to edit accounts (they can see them in the list)
    # But only superusers can change roles
    
    try:
        edit_user = User.objects.get(id=user_id)
    except User.DoesNotExist:
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': False, 'error': 'User not found'}, status=404)
        messages.error(request, 'User not found.')
        return redirect('app:list_accounts')
    
    try:
        profile = edit_user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=edit_user)
    
    if request.method == 'POST':
        # Update user fields
        edit_user.username = request.POST.get('username', edit_user.username)
        edit_user.email = request.POST.get('email', edit_user.email)
        edit_user.first_name = request.POST.get('first_name', edit_user.first_name)
        edit_user.last_name = request.POST.get('last_name', edit_user.last_name)
        edit_user.is_active = request.POST.get('is_active') == 'on'
        
        # Update role (only superusers can change roles)
        role = request.POST.get('role', 'user')
        if request.user.is_superuser:
            if role == 'admin':
                edit_user.is_superuser = True
                edit_user.is_staff = True
            elif role == 'staff':
                edit_user.is_superuser = False
                edit_user.is_staff = True
            elif role == 'user':
                edit_user.is_superuser = False
                edit_user.is_staff = False
        
        edit_user.save()
        
        # Update profile fields
        profile.phone = request.POST.get('phone', profile.phone)
        profile.country = request.POST.get('country', profile.country)
        if 'profile_photo' in request.FILES:
            profile.profile_photo = request.FILES['profile_photo']
        profile.save()
        
        # Return JSON for AJAX requests, otherwise redirect
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({'success': True})
        
        messages.success(request, 'Account updated successfully.')
        return redirect('app:view_account', user_id=user_id)
    
    return render(request, 'app/edit_user_account.html', {
        'edit_user': edit_user,
        'profile': profile,
    })

@login_required
def delete_user_account(request, user_id):
    """Delete user account"""
    if not (request.user.is_superuser or request.user.is_staff):
        return JsonResponse({'success': False, 'error': 'Permission denied'}, status=403)
    
    if request.user.id == user_id:
        return JsonResponse({'success': False, 'error': 'Cannot delete your own account'}, status=400)
    
    try:
        user = User.objects.get(id=user_id)
        user.delete()
        return JsonResponse({'success': True})
    except User.DoesNotExist:
        return JsonResponse({'success': False, 'error': 'User not found'}, status=404)
    except Exception as e:
        return JsonResponse({'success': False, 'error': str(e)}, status=500)

@login_required
def edit_account(request):
    """Edit Account page view"""
    user = request.user
    try:
        profile = user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=user)
    
    if request.method == 'POST':
        # Update user fields
        user.first_name = request.POST.get('first_name', '')
        user.last_name = request.POST.get('last_name', '')
        user.save()
        
        # Update profile fields
        profile.phone = request.POST.get('phone', '')
        profile.country = request.POST.get('country', '')
        if 'profile_photo' in request.FILES:
            profile.profile_photo = request.FILES['profile_photo']
        profile.save()
        
        messages.success(request, 'Account updated successfully!')
        return redirect('app:account_detail')
    
    return render(request, 'app/edit_account.html', {
        'user': user,
        'profile': profile,
        'active_page': 'edit_account',
    })

@login_required
def change_email(request):
    """Change Email page view"""
    user = request.user
    try:
        profile = user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=user)
    
    if request.method == 'POST':
        new_email = request.POST.get('new_email', '').strip()
        password = request.POST.get('password', '')
        
        # Verify password
        if not user.check_password(password):
            messages.error(request, 'Incorrect password. Please try again.')
        elif not new_email:
            messages.error(request, 'Please provide a new email address.')
        elif User.objects.filter(email=new_email).exclude(id=user.id).exists():
            messages.error(request, 'This email is already in use by another account.')
        else:
            user.email = new_email
            user.save()
            messages.success(request, 'Email updated successfully!')
            return redirect('app:account_detail')
    
    return render(request, 'app/change_email.html', {
        'user': user,
        'profile': profile,
        'active_page': 'change_email',
    })

@login_required
def reset_password(request):
    """Reset Password page view"""
    user = request.user
    try:
        profile = user.profile
    except UserProfile.DoesNotExist:
        profile = UserProfile.objects.create(user=user)
    
    if request.method == 'POST':
        old_password = request.POST.get('old_password', '')
        new_password = request.POST.get('new_password', '')
        confirm_password = request.POST.get('confirm_password', '')
        
        # Verify old password
        if not user.check_password(old_password):
            messages.error(request, 'Incorrect current password.')
        elif not new_password:
            messages.error(request, 'Please provide a new password.')
        elif len(new_password) < 8:
            messages.error(request, 'Password must be at least 8 characters long.')
        elif new_password != confirm_password:
            messages.error(request, 'New passwords do not match.')
        else:
            user.set_password(new_password)
            user.save()
            # Re-authenticate user after password change
            from django.contrib.auth import update_session_auth_hash
            update_session_auth_hash(request, user)
            messages.success(request, 'Password changed successfully!')
            return redirect('app:account_detail')
    
    return render(request, 'app/reset_password.html', {
        'user': user,
        'profile': profile,
        'active_page': 'reset_password',
    })
